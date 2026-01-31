"""
Word 文档处理 API
"""
from fastapi import APIRouter, UploadFile, File, HTTPException, Query
from fastapi.responses import FileResponse, HTMLResponse
from typing import Optional
from datetime import datetime
import html
from app.models.schemas import WordDocumentProcessResponse, DocumentEpisodesResponse, DocumentListResponse, DocumentListItem, DocumentContentResponse, RealtimeParsedContentResponse
from app.services.word_document_service import WordDocumentService
from app.services.version_migration_service import VersionMigrationService
from app.services.version_comparison_service import VersionComparisonService
from app.core.graphiti_client import get_graphiti_instance
from graphiti_core.nodes import EpisodicNode
import os
import uuid
import logging

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/word-document", tags=["Word文档处理"])


@router.post("/upload", response_model=WordDocumentProcessResponse)
async def upload_word_document(
    file: UploadFile = File(...),
    provider: str = Query("qianwen", description="LLM提供商"),
    max_tokens_per_section: int = Query(8000, ge=1000, le=20000, description="每个章节的最大token数")
):
    """
    上传并处理 Word 文档
    
    支持 .docx 和 .doc 格式
    """
    try:
        upload_dir = "uploads"
        os.makedirs(upload_dir, exist_ok=True)
        
        file_extension = os.path.splitext(file.filename)[1]
        if file_extension not in [".docx", ".doc"]:
            raise HTTPException(status_code=400, detail="只支持 .docx 和 .doc 格式")
        
        document_name = file.filename or f"文档_{uuid.uuid4()}"
        
        # 先保存文件到临时位置，然后处理文档获取正确的 group_id
        # 临时使用 UUID 作为文件名，处理完成后会重命名
        temp_document_id = str(uuid.uuid4())
        temp_file_path = os.path.join(upload_dir, f"{temp_document_id}{file_extension}")
        
        # 保存文件
        with open(temp_file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        logger.info(f"文件已保存到临时位置: {temp_file_path}")
        
        # 处理文档（会解析文档内容，生成正确的 group_id）
        result = await WordDocumentService.process_word_document(
            file_path=temp_file_path,
            document_name=document_name,
            provider=provider,
            max_tokens_per_section=max_tokens_per_section
        )
        
        # 处理完成后，从结果中获取 group_id，重命名文件
        group_id = result.get("document_id")  # process_word_document 返回的 document_id 就是 group_id
        if group_id:
            # 提取版本号用于文件名
            version, _ = WordDocumentService._extract_version(document_name)
            
            # 使用 group_id 作为文件名
            if version:
                final_file_name = f"{group_id}_{version}{file_extension}"
            else:
                final_file_name = f"{group_id}{file_extension}"
            
            final_file_path = os.path.join(upload_dir, final_file_name)
            
            # 如果文件已存在，添加时间戳
            if os.path.exists(final_file_path):
                from datetime import datetime
                timestamp = datetime.now().strftime('%H%M%S')
                final_file_name = f"{group_id}_{version}_{timestamp}{file_extension}" if version else f"{group_id}_{timestamp}{file_extension}"
                final_file_path = os.path.join(upload_dir, final_file_name)
            
            # 重命名文件
            os.rename(temp_file_path, final_file_path)
            logger.info(f"文件已重命名为: {final_file_path} (group_id={group_id})")
            
            # 更新 Neo4j 中的 file_path
            from app.core.neo4j_client import neo4j_client
            update_file_path_query = """
            MATCH (e:Episodic)
            WHERE e.group_id = $group_id
              AND e.name CONTAINS '文档概览'
            SET e.file_path = $file_path,
                e.original_filename = $original_filename
            RETURN e.uuid as uuid
            """
            neo4j_client.execute_write(update_file_path_query, {
                "group_id": group_id,
                "file_path": final_file_path,
                "original_filename": os.path.basename(final_file_path)
            })
            logger.info(f"已更新 Neo4j 中的文件路径: {final_file_path}")
        
        return WordDocumentProcessResponse(**result)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"处理 Word 文档失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")


@router.get("/{document_id}/episodes", response_model=DocumentEpisodesResponse)
async def get_document_episodes(
    document_id: str,
    include_content: bool = Query(False, description="是否包含 Episode 内容"),
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    获取文档的所有 Episode 信息
    """
    try:
        from app.core.neo4j_client import neo4j_client
        from app.core.utils import serialize_neo4j_properties
        
        # 查询文档级 Episode
        query = """
        MATCH (e:Episodic)
        WHERE e.group_id = $document_id
          AND e.name CONTAINS '文档概览'
        RETURN e.uuid as uuid, e.name as name, e.created_at as created_at, 
               e.version as version, e.version_number as version_number,
               e.document_name as document_name,
               properties(e) as properties
        ORDER BY e.created_at DESC
        LIMIT 1
        """
        
        result = neo4j_client.execute_query(query, {"document_id": document_id})
        
        document_episode = None
        if result:
            data = result[0]
            document_episode = {
                "uuid": data.get("uuid"),
                "name": data.get("name"),
                "created_at": data.get("created_at").isoformat() if data.get("created_at") and hasattr(data.get("created_at"), "isoformat") else str(data.get("created_at", "")),
                "version": data.get("version"),
                "version_number": data.get("version_number"),
                "document_name": data.get("document_name"),
                "properties": serialize_neo4j_properties(data.get("properties", {}))
            }
        
        # 查询章节级 Episode
        section_query = """
        MATCH (e:Episodic)
        WHERE e.group_id = $document_id
          AND e.name CONTAINS '章节_'
        RETURN e.uuid as uuid, e.name as name, e.created_at as created_at,
               e.version as version, e.version_number as version_number,
               properties(e) as properties
        ORDER BY e.name ASC
        """
        
        section_result = neo4j_client.execute_query(section_query, {"document_id": document_id})
        section_episodes = []
        for record in section_result:
            section_episodes.append({
                "uuid": record.get("uuid"),
                "name": record.get("name"),
                "created_at": record.get("created_at").isoformat() if record.get("created_at") and hasattr(record.get("created_at"), "isoformat") else str(record.get("created_at", "")),
                "version": record.get("version"),
                "version_number": record.get("version_number"),
                "properties": serialize_neo4j_properties(record.get("properties", {}))
            })
        
        # 查询图片 Episode
        image_query = """
        MATCH (e:Episodic)
        WHERE e.group_id = $document_id
          AND e.name CONTAINS '图片_'
        RETURN e.uuid as uuid, e.name as name, e.created_at as created_at,
               e.version as version, e.version_number as version_number,
               properties(e) as properties
        ORDER BY e.name ASC
        """
        
        image_result = neo4j_client.execute_query(image_query, {"document_id": document_id})
        image_episodes = []
        for record in image_result:
            image_episodes.append({
                "uuid": record.get("uuid"),
                "name": record.get("name"),
                "created_at": record.get("created_at").isoformat() if record.get("created_at") and hasattr(record.get("created_at"), "isoformat") else str(record.get("created_at", "")),
                "version": record.get("version"),
                "version_number": record.get("version_number"),
                "properties": serialize_neo4j_properties(record.get("properties", {}))
            })
        
        # 查询表格 Episode
        table_query = """
        MATCH (e:Episodic)
        WHERE e.group_id = $document_id
          AND e.name CONTAINS '表格_'
        RETURN e.uuid as uuid, e.name as name, e.created_at as created_at,
               e.version as version, e.version_number as version_number,
               properties(e) as properties
        ORDER BY e.name ASC
        """
        
        table_result = neo4j_client.execute_query(table_query, {"document_id": document_id})
        table_episodes = []
        for record in table_result:
            table_episodes.append({
                "uuid": record.get("uuid"),
                "name": record.get("name"),
                "created_at": record.get("created_at").isoformat() if record.get("created_at") and hasattr(record.get("created_at"), "isoformat") else str(record.get("created_at", "")),
                "version": record.get("version"),
                "version_number": record.get("version_number"),
                "properties": serialize_neo4j_properties(record.get("properties", {}))
            })
        
        return DocumentEpisodesResponse(
            document_id=document_id,
            document_episode=document_episode,
            section_episodes=section_episodes,
            image_episodes=image_episodes,
            table_episodes=table_episodes
        )
    except Exception as e:
        logger.error(f"获取文档 Episode 信息失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")


@router.get("/", response_model=DocumentListResponse)
async def list_documents(
    offset: int = Query(0, ge=0, description="偏移量"),
    limit: int = Query(20, ge=1, le=100, description="每页数量"),
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    获取文档列表
    """
    try:
        from app.core.neo4j_client import neo4j_client
        from app.core.utils import serialize_neo4j_properties
        
        logger.info(f"开始获取文档列表: offset={offset}, limit={limit}, provider={provider}")
        
        # 查询所有文档级 Episode（优先查找"文档概览"类型的Episode）
        query = """
        MATCH (e:Episodic)
        WHERE e.group_id STARTS WITH 'doc_'
          AND e.name CONTAINS '文档概览'
        WITH e.group_id as group_id,
             e.version as version,
             e.version_number as version_number,
             e.name as name,
             e.created_at as created_at,
             e.document_name as document_name
        ORDER BY e.group_id, e.version_number ASC
        RETURN DISTINCT group_id,
               version,
               version_number,
               name,
               created_at,
               document_name
        """
        
        logger.info("执行文档列表查询...")
        result = neo4j_client.execute_query(query, {})
        logger.info(f"查询返回 {len(result) if result else 0} 条记录")
        
        documents = []
        for record in result:
            try:
                group_id = record.get("group_id", "")
                if not group_id:
                    continue
                    
                document_name = record.get("name", "")
                created_at = record.get("created_at")
                version = record.get("version")
                version_number = record.get("version_number")
                document_name_full = record.get("document_name", document_name)
                
                # 统计各类 Episode 数量和总数
                stats_query = """
                MATCH (e:Episodic)
                WHERE e.group_id = $group_id
                  AND (e.version = $version OR (e.version IS NULL AND $version IS NULL))
                WITH e.name as name
                RETURN 
                  count(*) as total,
                  sum(CASE WHEN name CONTAINS '章节_' THEN 1 ELSE 0 END) as sections,
                  sum(CASE WHEN name CONTAINS '图片_' THEN 1 ELSE 0 END) as images,
                  sum(CASE WHEN name CONTAINS '表格_' THEN 1 ELSE 0 END) as tables
                """
                
                stats_result = neo4j_client.execute_query(stats_query, {
                    "group_id": group_id,
                    "version": version
                })
                episode_count = stats_result[0].get("total", 0) if stats_result else 0
                section_count = stats_result[0].get("sections", 0) if stats_result else 0
                image_count = stats_result[0].get("images", 0) if stats_result else 0
                table_count = stats_result[0].get("tables", 0) if stats_result else 0
            except Exception as e:
                logger.warning(f"处理文档记录失败: {e}", exc_info=True)
                continue
            
            documents.append(DocumentListItem(
                document_id=group_id,
                document_name=document_name_full or document_name,
                created_at=created_at.isoformat() if created_at and hasattr(created_at, "isoformat") else str(created_at or ""),
                total_episodes=episode_count,
                statistics={
                    "total_sections": section_count,
                    "total_images": image_count,
                    "total_tables": table_count,
                    "version": version,
                    "version_number": version_number
                }
            ))
        
        # 按创建时间降序排序
        documents.sort(key=lambda x: x.created_at or "", reverse=True)
        
        # 分页
        total = len(documents)
        documents = documents[offset:offset + limit]
        
        return DocumentListResponse(
            documents=documents,
            total=total
        )
    except Exception as e:
        logger.error(f"获取文档列表失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")


@router.get("/migration/analyze")
async def analyze_document_versions(
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    分析现有文档，识别同一需求的不同版本
    
    返回文档分组信息，用于迁移准备
    """
    try:
        doc_groups = VersionMigrationService.analyze_existing_documents()
        
        # 格式化返回结果
        result = {
            "total_groups": len(doc_groups),
            "document_groups": []
        }
        
        for base_group_id, versions in doc_groups.items():
            result["document_groups"].append({
                "base_group_id": base_group_id,
                "version_count": len(versions),
                "versions": [
                    {
                        "old_group_id": v["old_group_id"],
                        "version": v["version"],
                        "version_number": v["version_number"],
                        "document_name": v["document_name"],
                        "episode_count": v["episode_count"],
                        "created_at": v["created_at"].isoformat() if v.get("created_at") and hasattr(v["created_at"], "isoformat") else str(v.get("created_at", ""))
                    }
                    for v in versions
                ]
            })
        
        return result
    except Exception as e:
        logger.error(f"分析文档版本失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"分析失败: {str(e)}")


@router.post("/migration/migrate")
async def migrate_document_versions(
    dry_run: bool = Query(True, description="是否为试运行（不实际修改数据）"),
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    迁移所有文档到新的 group_id 格式
    
    Args:
        dry_run: 是否为试运行（默认 True，只分析不修改）
    
    Returns:
        迁移结果统计
    """
    try:
        result = VersionMigrationService.migrate_all_documents(dry_run=dry_run)
        
        return {
            "success": True,
            "dry_run": dry_run,
            "total_groups": result["total_groups"],
            "total_episodes_updated": result["total_episodes_updated"],
            "total_errors": result["total_errors"],
            "migrations": result["migrations"]
        }
    except Exception as e:
        logger.error(f"迁移文档版本失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"迁移失败: {str(e)}")


@router.get("/{document_id}/graph", response_model=dict)
async def get_document_graph(
    document_id: str,
    provider: str = Query("qianwen", description="LLM提供商"),
    limit: int = Query(500, ge=1, le=2000, description="返回节点数量限制")
):
    """
    获取文档相关的知识图谱数据
    
    返回该文档所有Episode相关的实体和关系
    """
    try:
        from app.core.neo4j_client import neo4j_client
        from app.core.utils import serialize_neo4j_properties
        
        # 查询该文档所有Episode相关的实体和关系
        # 优化：先限制节点数量，再查询关系，避免查询过多数据
        query = """
        // 第一步：找到该文档的所有Episode（通常数量不多，不需要限制）
        MATCH (e:Episodic)
        WHERE e.group_id = $document_id
        WITH collect(e) as episodes, collect(e.uuid) as episode_uuids
        
        // 第二步：查询这些Episode相关的实体（通过group_id），限制数量
        MATCH (n:Entity)
        WHERE n.group_id = $document_id
        WITH collect(DISTINCT n)[0..$limit] as entities, episodes, episode_uuids
        
        // 第三步：查询Community节点（如果group_id是字符串，检查是否包含；如果是列表，检查是否在列表中）
        // 使用OPTIONAL MATCH确保即使没有Community也能继续查询
        OPTIONAL MATCH (c:Community)
        WHERE (c.group_id = $document_id OR 
               (c.group_id IS NOT NULL AND 
                (toString(c.group_id) CONTAINS $document_id)))
        WITH coalesce(collect(DISTINCT c), []) as communities, entities, episodes, episode_uuids
        
        // 第四步：合并Episode节点、Entity节点和Community节点
        WITH entities + episodes + communities as all_nodes, episode_uuids
        
        // 第五步：查询这些节点之间的关系（包括Entity之间的关系、Episode与Entity之间的关系、Community与Entity之间的关系）
        // 只查询连接all_nodes中节点的关系，并限制关系数量
        OPTIONAL MATCH (a)-[r:RELATES_TO|MENTIONS|CONTAINS|HAS_MEMBER]->(b)
        WHERE (r.group_id = $document_id OR r.episode_uuid IN episode_uuids)
          AND a IN all_nodes 
          AND b IN all_nodes
        WITH all_nodes, coalesce(collect(DISTINCT r), [])[0..$limit] as relations
        
        RETURN 
          [node IN all_nodes | {
            id: id(node),
            labels: labels(node),
            properties: properties(node),
            content: node.content
          }] as nodes,
          [rel IN relations | {
            id: id(rel),
            source: id(startNode(rel)),
            target: id(endNode(rel)),
            type: type(rel),
            properties: properties(rel)
          }] as edges
        """
        
        result = neo4j_client.execute_query(query, {"document_id": document_id, "limit": limit})
        
        if not result:
            return {"nodes": [], "edges": []}
        
        data = result[0]
        
        # 处理节点
        nodes_dict = {}
        for node_data in data.get("nodes", []):
            if node_data.get("id") is not None:
                node_id = str(node_data["id"])
                props = node_data.get("properties", {})
                nodes_dict[node_id] = {
                    "id": node_id,
                    "labels": node_data.get("labels", []),
                    "name": props.get("name", ""),
                    "type": node_data.get("labels", ["Entity"])[0] if node_data.get("labels") else "Entity",
                    "properties": serialize_neo4j_properties(props),
                    "content": node_data.get("content")
                }
        
        # 处理边
        edges = []
        for edge_data in data.get("edges", []):
            if edge_data.get("id") is not None and edge_data.get("source") is not None:
                source_id = str(edge_data["source"])
                target_id = str(edge_data["target"])
                # 确保source和target节点都存在
                if source_id in nodes_dict and target_id in nodes_dict:
                    edges.append({
                        "id": str(edge_data["id"]),
                        "source": source_id,
                        "target": target_id,
                        "type": edge_data.get("type", ""),
                        "properties": serialize_neo4j_properties(edge_data.get("properties", {}))
                    })
        
        return {
            "nodes": list(nodes_dict.values()),
            "edges": edges,
            "document_id": document_id
        }
    except Exception as e:
        logger.error(f"获取文档图谱数据失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")


def _find_document_file_path(document_id: str) -> tuple[str, str]:
    """
    查找文档的文件路径
    
    Args:
        document_id: 文档ID（group_id）或 upload_{upload_id} 格式
    
    Returns:
        (file_path, original_filename)
    
    Raises:
        HTTPException: 如果找不到文件
    """
    # ========== 优先处理 upload_{upload_id} 格式（从MySQL查找） ==========
    if document_id.startswith("upload_"):
        try:
            upload_id_str = document_id.replace("upload_", "")
            upload_id = int(upload_id_str)
            
            from app.core.mysql_client import SessionLocal
            from app.models.document_upload import DocumentUpload
            
            db = SessionLocal()
            try:
                document = db.query(DocumentUpload).filter(DocumentUpload.id == upload_id).first()
                if document and document.file_path:
                    # 将相对路径转换为绝对路径
                    if not os.path.isabs(document.file_path):
                        file_path = os.path.join("/app", document.file_path)
                    else:
                        file_path = document.file_path
                    
                    if os.path.exists(file_path):
                        logger.info(f"从MySQL查询到文件路径: file_path={file_path}, upload_id={upload_id}")
                        return file_path, document.file_name or os.path.basename(file_path)
                    else:
                        logger.warning(f"MySQL中的文件路径不存在: {file_path}, upload_id={upload_id}")
                else:
                    logger.warning(f"MySQL中未找到upload_id={upload_id}的文档记录")
            finally:
                db.close()
        except (ValueError, Exception) as e:
            logger.warning(f"处理upload_{document_id}格式失败: {e}，继续尝试Neo4j查找")
    
    # ========== 从Neo4j查找（原有逻辑） ==========
    from app.core.neo4j_client import neo4j_client
    
    # 查询文档级 Episode 获取文件路径
    query = """
    MATCH (e:Episodic)
    WHERE e.group_id = $document_id
      AND e.name CONTAINS '文档概览'
    RETURN e.file_path as file_path, e.original_filename as original_filename,
           properties(e) as properties, e.group_id as group_id, e.version as version
    ORDER BY e.created_at DESC
    LIMIT 1
    """
    
    result = neo4j_client.execute_query(query, {"document_id": document_id})
    
    file_path = None
    original_filename = None
    group_id = None
    
    if result:
        data = result[0]
        props = data.get("properties", {})
        file_path = data.get("file_path") or props.get("file_path")
        original_filename = data.get("original_filename") or props.get("original_filename")
        group_id = data.get("group_id")
        version = data.get("version")
        logger.info(f"从Neo4j查询到文件路径: file_path={file_path}, document_id={document_id}, group_id={group_id}, version={version}")
    
    # 如果文件路径不存在，尝试查询同一group_id的所有版本，找到有file_path的版本
    if not file_path or not os.path.exists(file_path):
        if group_id:
            query_all_versions = """
            MATCH (e:Episodic)
            WHERE e.group_id = $group_id
              AND e.name CONTAINS '文档概览'
              AND e.file_path IS NOT NULL
              AND e.file_path <> ''
            RETURN e.file_path as file_path, e.original_filename as original_filename,
                   e.version as version
            ORDER BY e.created_at DESC
            LIMIT 1
            """
            result_all = neo4j_client.execute_query(query_all_versions, {"group_id": group_id})
            if result_all:
                data_all = result_all[0]
                file_path = data_all.get("file_path")
                original_filename = data_all.get("original_filename") or original_filename
                logger.info(f"从同一group_id的其他版本找到文件路径: file_path={file_path}, version={data_all.get('version')}")
    
    # 如果文件路径仍然不存在，尝试从uploads目录查找
    if not file_path or not os.path.exists(file_path):
        upload_dir = "uploads"
        if os.path.exists(upload_dir):
            import glob
            # 首先尝试精确匹配document_id
            possible_files = glob.glob(os.path.join(upload_dir, f"{document_id}.*"))
            if not possible_files and group_id:
                # 如果精确匹配失败，尝试使用group_id的基础部分匹配
                base_group_id = group_id.split('_V')[0] if '_V' in group_id else group_id
                possible_files = glob.glob(os.path.join(upload_dir, f"*{base_group_id}*"))
            if not possible_files:
                # 如果还是找不到，尝试部分匹配document_id
                possible_files = glob.glob(os.path.join(upload_dir, f"*{document_id}*"))
            
            # 如果还是找不到，记录警告但不要fallback到所有文件
            if not possible_files:
                logger.warning(f"无法找到document_id={document_id}对应的文件，Neo4j中的file_path={file_path}, group_id={group_id}")
                # 如果是 upload_{upload_id} 格式，提供更详细的错误信息
                if document_id.startswith("upload_"):
                    raise HTTPException(status_code=404, detail=f"文档文件不存在: document_id={document_id}（请确认文档已上传并解析）")
                else:
                    raise HTTPException(status_code=404, detail=f"文档文件不存在: document_id={document_id}")
            
            # 如果找到多个文件，选择第一个（应该只有一个）
            file_path = possible_files[0]
            logger.info(f"从uploads目录找到文件: {file_path} (document_id={document_id})")
    
    return file_path, original_filename


@router.get("/{document_id}/content", response_model=DocumentContentResponse)
async def get_document_content(
    document_id: str,
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    获取文档的原始内容和解析后的全部内容
    """
    try:
        from app.core.neo4j_client import neo4j_client
        from app.core.utils import serialize_neo4j_properties
        
        # 查找文件路径
        file_path, original_filename = _find_document_file_path(document_id)
        
        # 读取原始文档内容
        original_content = None
        if file_path and os.path.exists(file_path):
            try:
                from docx import Document
                doc = Document(file_path)
                original_content = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                logger.warning(f"读取原始文档失败: {e}")
        
        # 查询所有章节Episode，合并内容
        section_query = """
        MATCH (e:Episodic)
        WHERE e.group_id = $document_id
          AND e.name CONTAINS '章节_'
        RETURN e.content as content, e.name as name
        ORDER BY e.name ASC
        """
        
        section_result = neo4j_client.execute_query(section_query, {"document_id": document_id})
        
        parsed_content = ""
        for record in section_result:
            content = record.get("content", "")
            if content:
                parsed_content += content + "\n\n"
        
        return DocumentContentResponse(
            document_id=document_id,
            original_content=original_content,
            parsed_content=parsed_content.strip(),
            file_path=file_path,
            original_filename=original_filename
        )
    except Exception as e:
        logger.error(f"获取文档内容失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")


@router.get("/{document_id}/download")
async def download_document(
    document_id: str,
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    下载原始Word文档
    """
    try:
        # 查找文件路径
        file_path, original_filename = _find_document_file_path(document_id)
        
        return FileResponse(
            path=file_path,
            filename=original_filename or os.path.basename(file_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"下载文档失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"下载失败: {str(e)}")


@router.get("/{document_id}/realtime-parse")
async def get_realtime_parsed_content(
    document_id: str,
    provider: str = Query("qianwen", description="LLM提供商"),
    max_tokens_per_section: int = Query(8000, ge=1000, le=20000, description="每个章节的最大token数")
):
    """
    实时解析Word文档（从原始文档重新解析）
    """
    try:
        # 查找文件路径
        file_path, _ = _find_document_file_path(document_id)
        
        # 解析文档
        doc_data = WordDocumentService._parse_word_document(file_path, document_id)
        
        # 按章节分块
        sections = WordDocumentService._split_by_sections(
            doc_data["structured_content"],
            max_tokens=max_tokens_per_section
        )
        
        # 构建解析后的内容
        parsed_content = ""
        for idx, section in enumerate(sections):
            section_content = WordDocumentService._build_section_content(
                section, doc_data, idx, document_id
            )
            parsed_content += section_content + "\n\n"
        
        return RealtimeParsedContentResponse(
            document_id=document_id,
            parsed_content=parsed_content.strip(),
            sections=[
                {
                    "title": s["title"],
                    "content": s["content"][:200] + "..." if len(s["content"]) > 200 else s["content"]
                }
                for s in sections
            ],
            statistics={
                "total_sections": len(sections),
                "total_images": len(doc_data.get("images", [])),
                "total_tables": len(doc_data.get("tables", [])),
                "total_links": len(doc_data.get("links", [])),
                "total_ole_objects": len(doc_data.get("ole_objects", []))
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"实时解析文档失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"解析失败: {str(e)}")


@router.get("/{document_id}/images/{image_id}")
async def get_document_image(
    document_id: str,
    image_id: str,
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    获取文档中提取的图片
    """
    try:
        image_path = os.path.abspath(f"uploads/extracted_images/{document_id}/{image_id}.png")
        
        if not os.path.exists(image_path):
            # 尝试其他格式
            for ext in [".jpg", ".jpeg", ".gif", ".bmp"]:
                alt_path = os.path.abspath(f"uploads/extracted_images/{document_id}/{image_id}{ext}")
                if os.path.exists(alt_path):
                    image_path = alt_path
                    break
            else:
                raise HTTPException(status_code=404, detail="图片文件不存在")
        
        return FileResponse(
            path=image_path,
            media_type="image/png"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取图片失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")


@router.get("/{document_id}/preview")
async def preview_word_document(
    document_id: str,
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    预览Word文档（包含嵌入文件的链接）
    """
    try:
        # 查找文件路径
        file_path, original_filename = _find_document_file_path(document_id)
        
        # 使用与预览嵌入Word文档相同的逻辑
        from docx import Document
        html_escape = html
        
        doc = Document(file_path)
        
        # 生成HTML内容
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Word文档预览</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; background: #f5f5f5; }
                .container { background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); max-width: 1200px; margin: 0 auto; }
                h2 { color: #1890ff; margin-top: 0; }
                h3 { color: #333; margin-top: 30px; margin-bottom: 10px; }
                h4 { color: #666; margin-top: 20px; margin-bottom: 10px; }
                p { line-height: 1.6; margin: 10px 0; }
                div { line-height: 1.6; margin: 10px 0; }
                img { max-width: 100%; height: auto; margin: 10px 0; display: block; }
                table { border-collapse: collapse; width: 100%; margin: 20px 0; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f0f2f5; font-weight: bold; }
                ul, ol { margin: 10px 0; padding-left: 30px; }
                li { margin: 5px 0; }
            </style>
        </head>
        <body>
            <div class="container">
            <h2>Word文档预览: """ + html_escape.escape(original_filename or os.path.basename(file_path)) + """</h2>
        """
        
        # 按照文档的实际顺序遍历所有元素（段落、表格、OLE对象）
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn
        import zipfile
        import base64
        import glob
        
        # 建立段落位置到ole_id的映射（通过检查已提取的OLE对象文件）
        ole_dir = os.path.abspath(f"uploads/extracted_ole/{document_id}")
        ole_id_map = {}  # {OLE对象索引: ole_id}
        ole_id_to_file_name = {}  # {ole_id: file_name} 用于存储正确的文件名
        ole_counter = 0
        
        if os.path.exists(ole_dir):
            # 查找所有ole_*.ext格式的文件（排除临时文件）
            ole_files = [f for f in glob.glob(os.path.join(ole_dir, "ole_*.*")) 
                         if not os.path.basename(f).startswith("ole_temp_")]
            # 按文件名排序，确保顺序一致（ole_1, ole_2, ole_3...）
            ole_files.sort()
            logger.info(f"找到 {len(ole_files)} 个OLE对象文件: {[os.path.basename(f) for f in ole_files]}")
            for ole_file in ole_files:
                filename = os.path.basename(ole_file)
                # 提取ole_id（例如：ole_1.docx -> ole_1）
                if filename.startswith("ole_") and "." in filename:
                    ole_id = filename.split(".")[0]  # ole_1
                    # 提取数字部分（ole_1 -> 1）
                    try:
                        ole_num = int(ole_id.split("_")[1])
                        ole_counter += 1
                        # 使用计数器作为OLE对象索引的映射
                        ole_id_map[ole_counter] = ole_id
                        # 存储文件名（使用实际的文件名，而不是临时名称）
                        ole_id_to_file_name[ole_id] = filename
                    except:
                        ole_counter += 1
                        ole_id_map[ole_counter] = ole_id
                        ole_id_to_file_name[ole_id] = filename
        
        # 遍历文档body中的所有元素，保持原始顺序
        para_idx = 0  # 段落索引计数器
        ole_idx = 0  # OLE对象索引计数器
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # 段落
                para = Paragraph(element, doc)
                para_idx += 1
                
                # 先检查段落中是否有图片或OLE对象
                has_images = False
                has_ole_objects = False
                para_html = ""
                para_seen_r_ids = set()  # 用于段落内去重
                
                for run_idx, run in enumerate(para.runs):
                    # 先检查OLE对象（在文本之前，确保顺序）
                    # 方法1：检查OLEObject元素（Microsoft Office格式）
                    try:
                        ole_elements = run._element.xpath('.//o:OLEObject', namespaces={
                            'o': 'urn:schemas-microsoft-com:office:office'
                        })
                    except:
                        ole_elements = run._element.xpath('.//*[local-name()="OLEObject"]')
                    
                    # 方法2：检查Object元素
                    try:
                        object_elements = run._element.xpath('.//o:Object', namespaces={
                            'o': 'urn:schemas-microsoft-com:office:office'
                        })
                    except:
                        object_elements = run._element.xpath('.//*[local-name()="Object"]')
                    
                    # 方法3：检查WPS格式的嵌入对象
                    wps_ole_elements = []
                    try:
                        run_xml = run._element.xml
                        if run_xml:
                            import xml.etree.ElementTree as ET
                            has_excel_keywords = any(keyword in run_xml.lower() for keyword in ['excel', 'xls', 'xlsx', 'spreadsheet'])
                            has_word_keywords = any(keyword in run_xml.lower() for keyword in ['word', 'docx', 'doc'])
                            has_ppt_keywords = any(keyword in run_xml.lower() for keyword in ['powerpoint', 'ppt', 'pptx'])
                            
                            if has_excel_keywords or has_word_keywords or has_ppt_keywords:
                                root = ET.fromstring(run_xml)
                                for elem in root.iter():
                                    tag = elem.tag.lower()
                                    if 'excel' in tag or 'ole' in tag or 'object' in tag or 'embed' in tag:
                                        prog_id = elem.get('ProgId', '') or elem.get('progid', '')
                                        if not prog_id:
                                            for attr_name, attr_value in elem.attrib.items():
                                                attr_lower = attr_value.lower() if attr_value else ""
                                                if 'excel' in attr_lower or 'xls' in attr_lower or 'word' in attr_lower or 'docx' in attr_lower or 'doc' in attr_lower:
                                                    prog_id = attr_value
                                                    break
                                        
                                        if prog_id or 'excel' in tag or 'xls' in tag or 'word' in tag or 'docx' in tag or has_excel_keywords or has_word_keywords:
                                            r_id = None
                                            for attr_name, attr_value in elem.attrib.items():
                                                if 'id' in attr_name.lower() and attr_value:
                                                    try:
                                                        r_id = elem.get(qn('r:id')) or attr_value
                                                        break
                                                    except:
                                                        r_id = attr_value
                                                        break
                                            if r_id and r_id not in para_seen_r_ids:
                                                wps_ole_elements.append((elem, r_id))
                                                para_seen_r_ids.add(r_id)
                    except Exception as e:
                        logger.debug(f"检查WPS格式OLE对象失败: {e}")
                    
                    # 合并所有找到的OLE元素
                    all_ole_elements = []
                    for ole_element in ole_elements:
                        r_id = ole_element.get(qn('r:id'))
                        if r_id and r_id not in para_seen_r_ids:
                            all_ole_elements.append((ole_element, r_id))
                            para_seen_r_ids.add(r_id)
                    
                    for obj_element in object_elements:
                        r_id = obj_element.get(qn('r:id'))
                        if r_id and r_id not in para_seen_r_ids:
                            all_ole_elements.append((obj_element, r_id))
                            para_seen_r_ids.add(r_id)
                    
                    all_ole_elements.extend(wps_ole_elements)
                    
                    # 处理找到的OLE对象
                    for ole_element, r_id in all_ole_elements:
                        if r_id:
                            try:
                                rel = para.part.rels[r_id]
                                prog_id = ole_element.get('ProgId', '')
                                
                                # 获取文件名
                                file_name = ""
                                if hasattr(rel, 'target_ref'):
                                    file_name = rel.target_ref
                                else:
                                    file_name = str(rel.target)
                                
                                if not file_name:
                                    file_name = prog_id or "嵌入文档"
                                
                                # 根据ProgId判断文件类型
                                file_type = "嵌入对象"
                                if 'Excel' in prog_id or 'excel' in prog_id.lower():
                                    file_type = "Excel文件"
                                elif 'Word' in prog_id or 'word' in prog_id.lower():
                                    file_type = "Word文档"
                                
                                # 使用OLE对象出现顺序来匹配ole_id（按顺序）
                                ole_idx += 1
                                ole_id = ole_id_map.get(ole_idx, f"ole_{ole_idx}")
                                
                                # 从已保存的文件中获取正确的文件名
                                actual_file_name = ole_id_to_file_name.get(ole_id, file_name)
                                # 如果文件名是ole_1.xlsx这样的格式，尝试提取更友好的名称
                                if actual_file_name.startswith("ole_") and "." in actual_file_name:
                                    # 保持ole_id前缀，但使用实际的文件扩展名
                                    display_name = actual_file_name
                                else:
                                    display_name = actual_file_name
                                
                                # 构建预览和下载链接
                                preview_url = f"/api/word-document/{document_id}/ole/{ole_id}?view=preview"
                                download_url = f"/api/word-document/{document_id}/ole/{ole_id}?view=download"
                                
                                para_html += f'''
                                <div style="margin: 15px 0; padding: 10px; background-color: #f9f9f9; border-left: 3px solid #1890ff;">
                                    <strong>{html_escape.escape(display_name)}</strong> ({file_type})
                                    <br>
                                    <a href="{preview_url}" target="_blank" style="color: #1890ff; margin-right: 15px;">查看</a>
                                    <a href="{download_url}" download style="color: #1890ff;">下载</a>
                                </div>
                                '''
                                has_ole_objects = True
                            except Exception as e:
                                logger.warning(f"提取OLE对象失败: {e}", exc_info=True)
                    
                    # 检查是否有图片
                    blips = run._element.xpath('.//*[local-name()="blip"]')
                    
                    if blips:
                        for blip in blips:
                            r_embed = blip.get(qn('r:embed'))
                            if r_embed:
                                try:
                                    rel = para.part.rels[r_embed]
                                    if hasattr(rel, 'target_ref'):
                                        image_path = rel.target_ref
                                    else:
                                        image_path = str(rel.target)
                                    
                                    with zipfile.ZipFile(file_path, 'r') as zip_file:
                                        possible_paths = [
                                            image_path,
                                            f"word/{image_path}",
                                            image_path.replace('../', ''),
                                        ]
                                        
                                        image_data = None
                                        image_ext = '.png'
                                        
                                        for path in possible_paths:
                                            if path in zip_file.namelist():
                                                image_data = zip_file.read(path)
                                                if path.endswith('.jpg') or path.endswith('.jpeg'):
                                                    image_ext = '.jpg'
                                                    mime_type = 'image/jpeg'
                                                elif path.endswith('.png'):
                                                    image_ext = '.png'
                                                    mime_type = 'image/png'
                                                elif path.endswith('.gif'):
                                                    image_ext = '.gif'
                                                    mime_type = 'image/gif'
                                                else:
                                                    mime_type = 'image/png'
                                                
                                                image_base64 = base64.b64encode(image_data).decode('utf-8')
                                                para_html += f'<img src="data:{mime_type};base64,{image_base64}" style="max-width: 100%; height: auto; margin: 10px 0;" />'
                                                has_images = True
                                                break
                                except Exception as e:
                                    logger.debug(f"提取图片失败: {e}")
                    
                    # 添加文本内容（在OLE对象和图片之后）
                    run_text = run.text.strip()
                    if run_text:
                        para_html += html_escape.escape(run_text)
                
                # 如果段落有内容（文本、图片或OLE对象）
                if para_html or para.text.strip():
                    text = para.text.strip()
                    
                    # 检查是否是标题
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        try:
                            level_num = int(level)
                            html_content += f'<h{min(level_num + 2, 6)}>{html_escape.escape(text) if text else ""}</h{min(level_num + 2, 6)}>'
                        except:
                            html_content += f'<h3>{html_escape.escape(text) if text else ""}</h3>'
                        
                        # 如果有图片或OLE对象，添加到标题后
                        if has_images or has_ole_objects:
                            html_content += para_html
                    else:
                        # 普通段落
                        if has_images or has_ole_objects:
                            # 如果有图片或OLE对象，使用div包装
                            html_content += f'<div>{para_html}</div>'
                        elif text:
                            html_content += f'<p>{html_escape.escape(text)}</p>'
                        else:
                            html_content += '<p><br></p>'
                else:
                    html_content += '<p><br></p>'
            
            elif isinstance(element, CT_Tbl):
                # 表格 - 在文档中的实际位置显示
                table = Table(element, doc)
                html_content += '<table>'
                # 表头
                if table.rows:
                    html_content += '<thead><tr>'
                    for cell in table.rows[0].cells:
                        html_content += f'<th>{html_escape.escape(cell.text)}</th>'
                    html_content += '</tr></thead>'
                # 表格内容
                html_content += '<tbody>'
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0: # 跳过表头
                        continue
                    html_content += '<tr>'
                    for cell in row.cells:
                        html_content += f'<td>{html_escape.escape(cell.text)}</td>'
                    html_content += '</tr>'
                html_content += '</tbody>'
                html_content += '</table>'
        
        html_content += """
            </div>
        </body>
        </html>
        """
        
        return HTMLResponse(
            content=html_content,
            headers={
                "Content-Type": "text/html; charset=utf-8",
                "X-Content-Type-Options": "nosniff"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"预览Word文档失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"预览失败: {str(e)}")


@router.get("/{document_id}/ole/{ole_id}")
async def get_document_ole(
    document_id: str,
    ole_id: str,
    view: str = Query("download", description="查看模式：download 或 preview"),
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    获取文档中提取的嵌入文档（OLE对象）
    
    Args:
        view: "download" 下载文件, "preview" 在线预览（支持Excel和Word）
    """
    try:
        ole_dir = os.path.abspath(f"uploads/extracted_ole/{document_id}")
        
        # 查找文件（可能包含ole_前缀）
        ole_files = []
        if os.path.exists(ole_dir):
            import glob
            # 尝试多种可能的文件名格式
            # 注意：ole_id可能已经是"ole_1"格式，也可能只是"1"
            patterns = [
                f"{ole_id}.*",  # ole_1.* 或 1.*
                f"ole_{ole_id}.*",  # ole_ole_1.* 或 ole_1.*
                f"{ole_id.replace('ole_', '')}.*"  # 1.* 或 ole_1.*
            ]
            # 去重
            patterns = list(set(patterns))
            logger.debug(f"查找OLE文件模式: ole_id={ole_id}, patterns={patterns}")
            for pattern in patterns:
                ole_files.extend(glob.glob(os.path.join(ole_dir, pattern)))
        
        if not ole_files:
            raise HTTPException(status_code=404, detail="嵌入文档文件不存在")
        
        # 优先选择标准格式文件（.xlsx > .xls > .doc > .docx > .ppt > .pptx > 其他）
        preferred_extensions = ['.xlsx', '.xls', '.docx', '.doc', '.pptx', '.ppt']
        ole_path = None
        
        # 首先尝试找到标准格式文件
        for ext in preferred_extensions:
            for file_path in ole_files:
                if file_path.endswith(ext):
                    ole_path = file_path
                    break
            if ole_path:
                break
        
        # 如果没有找到标准格式，使用第一个文件
        if not ole_path:
            ole_path = ole_files[0]
        
        filename = os.path.basename(ole_path)
        logger.info(f"[{view.upper()}] 选择文件: {filename} (从 {len(ole_files)} 个文件中选择), ole_id={ole_id}, document_id={document_id}")
        
        # 如果是预览模式
        if view == "preview":
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext in [".xlsx", ".xls"]:
                try:
                    import pandas as pd
                    html_escape = html
                    
                    # 文件已经在保存时提取并转换为标准格式，直接读取即可
                    excel_path = ole_path
                    
                    # 使用pandas读取Excel文件（支持.xls和.xlsx）
                    # 首先验证文件格式（检查文件头）
                    with open(excel_path, 'rb') as f:
                        file_header = f.read(8)
                    
                    # 检查是否是ZIP格式（.xlsx）
                    is_zip = file_header[:2] == b'PK'
                    # 检查是否是OLE2格式（.xls）
                    is_ole2 = file_header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'
                    
                    logger.info(f"文件格式检查: 路径={excel_path}, ZIP={is_zip}, OLE2={is_ole2}, 扩展名={file_ext}")
                    
                    # 如果文件扩展名是.xlsx但内容是OLE2，说明提取有问题
                    if file_ext == '.xlsx' and is_ole2:
                        logger.error(f"文件扩展名是.xlsx但内容是OLE2格式，提取可能失败: {excel_path}")
                        raise HTTPException(
                            status_code=500,
                            detail=f"文件格式错误：文件扩展名是.xlsx但内容是OLE2格式。这可能是提取过程中的问题。错误信息: Can't find workbook in OLE2 compound document"
                        )
                    
                    # 根据扩展名选择引擎（.xlsx使用openpyxl，.xls使用xlrd）
                    try:
                        if file_ext == '.xls':
                            engine = 'xlrd'
                        elif file_ext == '.xlsx':
                            engine = 'openpyxl'
                        else:
                            # 如果扩展名不是标准的，根据文件内容判断
                            if is_zip:
                                engine = 'openpyxl'
                            elif is_ole2:
                                engine = 'xlrd'
                            else:
                                engine = 'openpyxl'  # 默认使用openpyxl
                        
                        logger.info(f"使用引擎: {engine} (扩展名: {file_ext}, ZIP: {is_zip}, OLE2: {is_ole2})")
                        excel_file = pd.ExcelFile(excel_path, engine=engine)
                    except Exception as e:
                        logger.warning(f"使用{engine}引擎读取失败，尝试自动检测: {e}")
                        # 如果指定引擎失败，让pandas自动检测
                        try:
                            excel_file = pd.ExcelFile(excel_path)
                        except Exception as e2:
                            logger.error(f"pandas自动检测也失败: {e2}")
                            raise HTTPException(
                                status_code=500,
                                detail=f"无法读取Excel文件。文件路径: {excel_path}, 扩展名: {file_ext}, ZIP格式: {is_zip}, OLE2格式: {is_ole2}, 错误: {str(e2)}"
                            )
                    
                    # 生成HTML表格
                    html_content = """
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <title>Excel预览</title>
                        <style>
                            body { font-family: Arial, sans-serif; margin: 20px; background: #f5f5f5; }
                            .container { background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }
                            h2 { color: #1890ff; margin-top: 0; }
                            h3 { color: #666; margin-top: 30px; }
                            table { border-collapse: collapse; width: 100%; margin: 20px 0; background: white; max-height: 80vh; overflow-y: auto; display: block; }
                            thead { display: table-header-group; }
                            tbody { display: table-row-group; max-height: 70vh; overflow-y: auto; display: block; }
                            tr { display: table-row; }
                            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; display: table-cell; }
                            th { background-color: #f0f2f5; font-weight: bold; color: #333; position: sticky; top: 0; z-index: 10; }
                            tr:nth-child(even) { background-color: #fafafa; }
                            tr:hover { background-color: #f0f7ff; }
                        </style>
                    </head>
                    <body>
                        <div class="container">
                        <h2>Excel文件预览: """ + html_escape.escape(filename) + """</h2>
                    """
                    
                    # 遍历所有工作表
                    for sheet_name in excel_file.sheet_names:
                        try:
                            # 读取工作表（显示全部数据）
                            # 对于.xls文件使用xlrd引擎，对于.xlsx文件使用openpyxl引擎
                            engine = 'xlrd' if file_ext == '.xls' else 'openpyxl'
                            df = pd.read_excel(excel_path, sheet_name=sheet_name, engine=engine)
                            
                            html_content += f'<h3>工作表: {html_escape.escape(sheet_name)} ({len(df)} 行, {len(df.columns)} 列)</h3>'
                            
                            if df.empty:
                                html_content += '<p>此工作表为空</p>'
                            else:
                                html_content += '<table>'
                                
                                # 表头
                                html_content += '<thead><tr>'
                                for col in df.columns:  # 显示所有列
                                    html_content += f'<th>{html_escape.escape(str(col) if pd.notna(col) else "")}</th>'
                                html_content += '</tr></thead><tbody>'
                                
                                # 数据行（显示所有行）
                                for idx, row in df.iterrows():
                                    html_content += '<tr>'
                                    for col in df.columns:
                                        cell_value = row[col]
                                        if pd.isna(cell_value):
                                            cell_value = ""
                                        html_content += f'<td>{html_escape.escape(str(cell_value))}</td>'
                                    html_content += '</tr>'
                                
                                html_content += '</tbody></table>'
                        except Exception as e:
                            logger.warning(f"读取工作表 {sheet_name} 失败: {e}")
                            html_content += f'<p style="color: red;">读取工作表 "{html_escape.escape(sheet_name)}" 失败: {html_escape.escape(str(e))}</p>'
                    
                    html_content += f'''
                        </div>
                    </body>
                    </html>
                    '''
                    
                    return HTMLResponse(
                        content=html_content,
                        headers={
                            "Content-Type": "text/html; charset=utf-8",
                            "X-Content-Type-Options": "nosniff"
                        }
                    )
                except Exception as e:
                    logger.error(f"Excel预览失败: {e}", exc_info=True)
                    # 如果预览失败，返回错误页面而不是直接下载
                    error_html = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <title>预览失败</title>
                        <style>
                            body {{ font-family: Arial, sans-serif; margin: 40px; background: #f5f5f5; }}
                            .container {{ background: white; padding: 40px; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center; }}
                            h2 {{ color: #ff4d4f; margin-top: 0; }}
                            p {{ color: #666; margin: 20px 0; }}
                            .download-btn {{ display: inline-block; margin-top: 20px; padding: 10px 20px; background: #1890ff; color: white; text-decoration: none; border-radius: 4px; }}
                        </style>
                    </head>
                    <body>
                        <div class="container">
                            <h2>预览加载失败</h2>
                            <p>无法预览此文件，请下载后查看。</p>
                            <p style="color: #999; font-size: 12px;">错误信息: {html.escape(str(e))}</p>
                        </div>
                    </body>
                    </html>
                    """
                    return HTMLResponse(
                        content=error_html,
                        headers={
                            "Content-Type": "text/html; charset=utf-8",
                            "X-Content-Type-Options": "nosniff"
                        }
                    )
            
            # Word文件预览
            elif file_ext in [".docx", ".doc"]:
                try:
                    from docx import Document
                    html_escape = html
                    
                    # 读取Word文档（仅支持.docx格式）
                    if file_ext == ".docx":
                        doc = Document(ole_path)
                    else:
                        # .doc格式需要特殊处理，暂时提示不支持
                        raise HTTPException(
                            status_code=400,
                            detail="暂不支持预览 .doc 格式文件，请下载后查看。支持预览 .docx 格式文件。"
                        )
                    
                    # 生成HTML内容
                    html_content = """
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <title>Word文档预览</title>
                        <style>
                            body { font-family: Arial, sans-serif; margin: 20px; background: #f5f5f5; }
                            .container { background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); max-width: 1200px; margin: 0 auto; }
                            h2 { color: #1890ff; margin-top: 0; }
                            h3 { color: #333; margin-top: 30px; margin-bottom: 10px; }
                            h4 { color: #666; margin-top: 20px; margin-bottom: 10px; }
                            p { line-height: 1.6; margin: 10px 0; }
                            div { line-height: 1.6; margin: 10px 0; }
                            img { max-width: 100%; height: auto; margin: 10px 0; display: block; }
                            table { border-collapse: collapse; width: 100%; margin: 20px 0; }
                            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                            th { background-color: #f0f2f5; font-weight: bold; }
                            ul, ol { margin: 10px 0; padding-left: 30px; }
                            li { margin: 5px 0; }
                        </style>
                    </head>
                    <body>
                        <div class="container">
                        <h2>Word文档预览: """ + html_escape.escape(filename) + """</h2>
                    """
                    
                    # 按照文档的实际顺序遍历所有元素（段落、表格、OLE对象）
                    from docx.oxml.text.paragraph import CT_P
                    from docx.oxml.table import CT_Tbl
                    from docx.table import Table
                    from docx.text.paragraph import Paragraph
                    from docx.oxml.ns import qn
                    import zipfile
                    import base64
                    import glob
                    
                    # 建立段落位置到ole_id的映射（通过检查已提取的OLE对象文件）
                    # 注意：预览嵌入的Word文档时，我们需要查找原始文档目录中的OLE对象文件
                    # 因为嵌入文件是保存在原始文档的目录中的
                    ole_dir = os.path.abspath(f"uploads/extracted_ole/{document_id}")
                    ole_id_map = {}  # {OLE对象索引: ole_id}
                    ole_counter = 0
                    
                    if os.path.exists(ole_dir):
                        # 查找所有ole_*.ext格式的文件（排除临时文件）
                        ole_files = [f for f in glob.glob(os.path.join(ole_dir, "ole_*.*")) 
                                     if not os.path.basename(f).startswith("ole_temp_")]
                        # 按文件名排序，确保顺序一致（ole_1, ole_2, ole_3...）
                        ole_files.sort()
                        logger.info(f"找到 {len(ole_files)} 个OLE对象文件: {[os.path.basename(f) for f in ole_files]}")
                        for ole_file in ole_files:
                            filename = os.path.basename(ole_file)
                            # 提取ole_id（例如：ole_1.docx -> ole_1）
                            if filename.startswith("ole_") and "." in filename:
                                ole_id = filename.split(".")[0]  # ole_1
                                # 提取数字部分（ole_1 -> 1）
                                try:
                                    ole_num = int(ole_id.split("_")[1])
                                    ole_counter += 1
                                    # 使用计数器作为OLE对象索引的映射
                                    ole_id_map[ole_counter] = ole_id
                                    logger.debug(f"OLE对象映射: {ole_counter} -> {ole_id} (文件: {filename})")
                                except:
                                    # 如果不是标准格式，使用计数器
                                    ole_counter += 1
                                    ole_id_map[ole_counter] = ole_id
                    
                    # 遍历文档body中的所有元素，保持原始顺序
                    # 注意：这里遍历的是嵌入的Word文档（ole_1.docx）的元素
                    # 如果这个嵌入文档中也有嵌入文件，它们应该已经被提取并保存在原始文档的目录中
                    para_idx = 0  # 段落索引计数器（相对于当前嵌入文档）
                    ole_idx = 0  # OLE对象索引计数器（相对于当前嵌入文档中的OLE对象）
                    for element in doc.element.body:
                        if isinstance(element, CT_P):
                            # 段落
                            para = Paragraph(element, doc)
                            para_idx += 1
                            
                            # 先检查段落中是否有图片或OLE对象
                            has_images = False
                            has_ole_objects = False
                            para_html = ""
                            
                            # 调试：记录段落信息
                            if para_idx in [26, 317]:
                                try:
                                    para_xml = para._element.xml[:1000] if para._element.xml else "None"
                                    logger.info(f"检查段落 {para_idx}: 文本长度={len(para.text)}, runs数量={len(para.runs)}, 段落XML预览={para_xml}")
                                except:
                                    logger.info(f"检查段落 {para_idx}: 文本长度={len(para.text)}, runs数量={len(para.runs)}")
                            
                            # 遍历段落中的runs，在一个循环中同时处理图片、文本和OLE对象，确保顺序正确
                            para_seen_r_ids = set()  # 用于段落内去重
                            for run_idx, run in enumerate(para.runs):
                                # 调试：记录关键段落的run XML内容
                                if para_idx in [26, 317]:
                                    try:
                                        run_xml_full = run._element.xml if run._element.xml else "None"
                                        # 检查XML中是否包含OLE相关关键词
                                        has_ole_keywords = any(kw in run_xml_full.lower() for kw in ['ole', 'object', 'embed', 'excel', 'word', 'docx'])
                                        logger.info(f"段落 {para_idx} run {run_idx}: XML长度={len(run_xml_full)}, 包含OLE关键词={has_ole_keywords}")
                                        if has_ole_keywords:
                                            logger.info(f"段落 {para_idx} run {run_idx}: XML内容={run_xml_full}")
                                    except Exception as e:
                                        logger.warning(f"获取段落 {para_idx} run {run_idx} XML失败: {e}")
                                
                                # 先检查OLE对象（在文本之前，确保顺序）
                                # 方法1：检查OLEObject元素（Microsoft Office格式）
                                try:
                                    ole_elements = run._element.xpath('.//o:OLEObject', namespaces={
                                        'o': 'urn:schemas-microsoft-com:office:office'
                                    })
                                except:
                                    ole_elements = run._element.xpath('.//*[local-name()="OLEObject"]')
                                
                                if para_idx in [26, 317]:
                                    logger.info(f"段落 {para_idx} run {run_idx}: OLEObject元素数量={len(ole_elements)}")
                                
                                # 方法2：检查Object元素
                                try:
                                    object_elements = run._element.xpath('.//o:Object', namespaces={
                                        'o': 'urn:schemas-microsoft-com:office:office'
                                    })
                                except:
                                    object_elements = run._element.xpath('.//*[local-name()="Object"]')
                                
                                # 方法3：检查WPS格式的嵌入对象（通过XML内容检查）
                                # 使用与解析代码相同的逻辑
                                wps_ole_elements = []
                                try:
                                    run_xml = run._element.xml
                                    if run_xml:
                                        import xml.etree.ElementTree as ET
                                        # 检查是否包含Excel、Word等关键词（与解析代码一致，但需要扩展以支持Word）
                                        # 注意：解析代码中只检查'excel', 'xls', 'xlsx', 'spreadsheet'，但我们需要也检查Word
                                        has_excel_keywords = any(keyword in run_xml.lower() for keyword in ['excel', 'xls', 'xlsx', 'spreadsheet'])
                                        has_word_keywords = any(keyword in run_xml.lower() for keyword in ['word', 'docx', 'doc'])
                                        has_ppt_keywords = any(keyword in run_xml.lower() for keyword in ['powerpoint', 'ppt', 'pptx'])
                                        
                                        if has_excel_keywords or has_word_keywords or has_ppt_keywords:
                                            root = ET.fromstring(run_xml)
                                            for elem in root.iter():
                                                tag = elem.tag.lower()
                                                # 检查是否包含OLE相关标签
                                                if 'excel' in tag or 'ole' in tag or 'object' in tag or 'embed' in tag:
                                                    # 提取ProgId或类型信息
                                                    prog_id = elem.get('ProgId', '') or elem.get('progid', '')
                                                    if not prog_id:
                                                        # 尝试从属性中获取
                                                        for attr_name, attr_value in elem.attrib.items():
                                                            attr_lower = attr_value.lower() if attr_value else ""
                                                            if 'excel' in attr_lower or 'xls' in attr_lower or 'word' in attr_lower or 'docx' in attr_lower or 'doc' in attr_lower:
                                                                prog_id = attr_value
                                                                break
                                                    
                                                    # 检查条件：如果有prog_id，或者标签中包含excel/xls/word等关键词
                                                    # 或者XML中包含相关关键词
                                                    if prog_id or 'excel' in tag or 'xls' in tag or 'word' in tag or 'docx' in tag or has_excel_keywords or has_word_keywords:
                                                        # 获取关系ID
                                                        r_id = None
                                                        for attr_name, attr_value in elem.attrib.items():
                                                            if 'id' in attr_name.lower() and attr_value:
                                                                try:
                                                                    r_id = elem.get(qn('r:id')) or attr_value
                                                                    break
                                                                except:
                                                                    r_id = attr_value
                                                                    break
                                                        if r_id and r_id not in para_seen_r_ids:
                                                            wps_ole_elements.append((elem, r_id))
                                                            para_seen_r_ids.add(r_id)
                                                            if para_idx in [26, 317]:
                                                                logger.info(f"段落 {para_idx}: 找到WPS格式OLE对象, r_id={r_id}, tag={tag}, prog_id={prog_id}")
                                except Exception as e:
                                    if para_idx in [26, 317]:
                                        logger.warning(f"检查WPS格式OLE对象失败: {e}", exc_info=True)
                                
                                # 合并所有找到的OLE元素
                                all_ole_elements = []
                                for ole_element in ole_elements:
                                    r_id = ole_element.get(qn('r:id'))
                                    if r_id and r_id not in para_seen_r_ids:
                                        all_ole_elements.append((ole_element, r_id))
                                        para_seen_r_ids.add(r_id)
                                        if para_idx in [26, 317]:
                                            logger.info(f"段落 {para_idx}: 找到OLEObject元素, r_id={r_id}")
                                
                                for obj_element in object_elements:
                                    r_id = obj_element.get(qn('r:id'))
                                    if r_id and r_id not in para_seen_r_ids:
                                        all_ole_elements.append((obj_element, r_id))
                                        para_seen_r_ids.add(r_id)
                                        if para_idx in [26, 317]:
                                            logger.info(f"段落 {para_idx}: 找到Object元素, r_id={r_id}")
                                
                                all_ole_elements.extend(wps_ole_elements)
                                
                                # 处理找到的OLE对象
                                for ole_element, r_id in all_ole_elements:
                                    if r_id:
                                        try:
                                            rel = para.part.rels[r_id]
                                            prog_id = ole_element.get('ProgId', '')
                                            
                                            # 获取文件名
                                            file_name = ""
                                            if hasattr(rel, 'target_ref'):
                                                file_name = rel.target_ref
                                            else:
                                                file_name = str(rel.target)
                                            
                                            if not file_name:
                                                file_name = prog_id or "嵌入文档"
                                            
                                            # 根据ProgId判断文件类型
                                            file_type = "嵌入对象"
                                            if 'Excel' in prog_id or 'excel' in prog_id.lower():
                                                file_type = "Excel文件"
                                            elif 'Word' in prog_id or 'word' in prog_id.lower():
                                                file_type = "Word文档"
                                            
                                            # 使用OLE对象出现顺序来匹配ole_id（按顺序）
                                            ole_idx += 1
                                            ole_id = ole_id_map.get(ole_idx, f"ole_{ole_idx}")
                                            logger.info(f"段落 {para_idx} 中的OLE对象 {ole_idx}: r_id={r_id}, prog_id={prog_id}, file_name={file_name}, ole_id={ole_id}")
                                            
                                            # 构建预览和下载链接
                                            preview_url = f"/api/word-document/{document_id}/ole/{ole_id}?view=preview"
                                            download_url = f"/api/word-document/{document_id}/ole/{ole_id}?view=download"
                                            
                                            para_html += f'''
                                            <div style="margin: 15px 0; padding: 10px; background-color: #f9f9f9; border-left: 3px solid #1890ff;">
                                                <strong>{html_escape.escape(file_name)}</strong> ({file_type})
                                                <br>
                                                <a href="{preview_url}" target="_blank" style="color: #1890ff; margin-right: 15px;">查看</a>
                                                <a href="{download_url}" download style="color: #1890ff;">下载</a>
                                            </div>
                                            '''
                                            has_ole_objects = True
                                        except Exception as e:
                                            logger.warning(f"提取OLE对象失败: {e}", exc_info=True)
                                
                                # 检查是否有图片
                                blips = run._element.xpath('.//*[local-name()="blip"]')
                                
                                if blips:
                                    for blip in blips:
                                        r_embed = blip.get(qn('r:embed'))
                                        if r_embed:
                                            try:
                                                rel = para.part.rels[r_embed]
                                                if hasattr(rel, 'target_ref'):
                                                    image_path = rel.target_ref
                                                else:
                                                    image_path = str(rel.target)
                                                
                                                with zipfile.ZipFile(ole_path, 'r') as zip_file:
                                                    possible_paths = [
                                                        image_path,
                                                        f"word/{image_path}",
                                                        image_path.replace('../', ''),
                                                    ]
                                                    
                                                    image_data = None
                                                    image_ext = '.png'
                                                    
                                                    for path in possible_paths:
                                                        if path in zip_file.namelist():
                                                            image_data = zip_file.read(path)
                                                            if path.endswith('.jpg') or path.endswith('.jpeg'):
                                                                image_ext = '.jpg'
                                                                mime_type = 'image/jpeg'
                                                            elif path.endswith('.png'):
                                                                image_ext = '.png'
                                                                mime_type = 'image/png'
                                                            elif path.endswith('.gif'):
                                                                image_ext = '.gif'
                                                                mime_type = 'image/gif'
                                                            else:
                                                                mime_type = 'image/png'
                                                            
                                                            image_base64 = base64.b64encode(image_data).decode('utf-8')
                                                            para_html += f'<img src="data:{mime_type};base64,{image_base64}" style="max-width: 100%; height: auto; margin: 10px 0;" />'
                                                            has_images = True
                                                            break
                                            except Exception as e:
                                                logger.debug(f"提取图片失败: {e}")
                                
                                # 添加文本内容（在OLE对象和图片之后）
                                run_text = run.text.strip()
                                if run_text:
                                    para_html += html_escape.escape(run_text)
                            
                            # 如果段落有内容（文本、图片或OLE对象）
                            if para_html or para.text.strip():
                                text = para.text.strip()
                                
                                # 检查是否是标题
                                if para.style.name.startswith('Heading'):
                                    level = para.style.name.replace('Heading ', '')
                                    try:
                                        level_num = int(level)
                                        html_content += f'<h{min(level_num + 2, 6)}>{html_escape.escape(text) if text else ""}</h{min(level_num + 2, 6)}>'
                                    except:
                                        html_content += f'<h3>{html_escape.escape(text) if text else ""}</h3>'
                                    
                                    # 如果有图片或OLE对象，添加到标题后
                                    if has_images or has_ole_objects:
                                        html_content += para_html
                                else:
                                    # 普通段落
                                    if has_images or has_ole_objects:
                                        # 如果有图片或OLE对象，使用div包装
                                        html_content += f'<div>{para_html}</div>'
                                    elif text:
                                        html_content += f'<p>{html_escape.escape(text)}</p>'
                                    else:
                                        html_content += '<p><br></p>'
                            else:
                                html_content += '<p><br></p>'
                        
                        elif isinstance(element, CT_Tbl):
                            # 表格 - 在文档中的实际位置显示
                            table = Table(element, doc)
                            html_content += '<table style="border-collapse: collapse; width: 100%; margin: 20px 0;">'
                            
                            # 处理表头（第一行）
                            if table.rows:
                                html_content += '<thead>'
                                html_content += '<tr>'
                                for cell in table.rows[0].cells:
                                    html_content += f'<th style="border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #f0f2f5; font-weight: bold;">{html_escape.escape(cell.text)}</th>'
                                html_content += '</tr>'
                                html_content += '</thead>'
                            
                            # 处理数据行
                            if len(table.rows) > 1:
                                html_content += '<tbody>'
                                for row in table.rows[1:]:
                                    html_content += '<tr>'
                                    for cell in row.cells:
                                        html_content += f'<td style="border: 1px solid #ddd; padding: 8px; text-align: left;">{html_escape.escape(cell.text)}</td>'
                                    html_content += '</tr>'
                                html_content += '</tbody>'
                            
                            html_content += '</table>'
                    
                    html_content += """
                        </div>
                    </body>
                    </html>
                    """
                    
                    return HTMLResponse(
                        content=html_content,
                        headers={
                            "Content-Type": "text/html; charset=utf-8",
                            "X-Content-Type-Options": "nosniff"
                        }
                    )
                except HTTPException:
                    raise
                except Exception as e:
                    logger.error(f"Word预览失败: {e}", exc_info=True)
                    error_html = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <title>预览失败</title>
                        <style>
                            body {{ font-family: Arial, sans-serif; margin: 40px; background: #f5f5f5; }}
                            .container {{ background: white; padding: 40px; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center; }}
                            h2 {{ color: #ff4d4f; margin-bottom: 16px; }}
                            p {{ color: #666; margin: 20px 0; }}
                        </style>
                    </head>
                    <body>
                        <div class="container">
                            <h2>预览加载失败</h2>
                            <p>无法预览此Word文件,请下载后查看。</p>
                            <p style="color: #999; font-size: 12px;">错误信息: {html.escape(str(e))}</p>
                        </div>
                    </body>
                    </html>
                    """
                    return HTMLResponse(
                        content=error_html,
                        headers={
                            "Content-Type": "text/html; charset=utf-8",
                            "X-Content-Type-Options": "nosniff"
                        }
                    )
            
            # 其他格式不支持预览
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"不支持预览 {file_ext} 格式的文件。目前仅支持 Excel (.xlsx, .xls) 和 Word (.docx) 格式的预览。"
                )
        
        # 下载模式
        return FileResponse(
            path=ole_path,
            filename=filename,
            media_type="application/octet-stream"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取嵌入文档失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")


@router.get("/{base_document_id}/versions")
async def get_document_versions(
    base_document_id: str,
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    查询同一需求的所有版本
    
    Args:
        base_document_id: 基础文档ID（group_id）
    
    Returns:
        版本列表，按版本号排序
    """
    try:
        versions = VersionComparisonService.get_document_versions(base_document_id)
        
        return {
            "base_document_id": base_document_id,
            "version_count": len(versions),
            "versions": versions
        }
    except Exception as e:
        logger.error(f"查询文档版本失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"查询失败: {str(e)}")


@router.get("/{base_document_id}/versions/{version}")
async def get_document_version(
    base_document_id: str,
    version: str,
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    查询特定版本的详细信息
    
    Args:
        base_document_id: 基础文档ID（group_id）
        version: 版本号（如 "V1", "V2"）
    
    Returns:
        版本详细信息
    """
    try:
        versions = VersionComparisonService.get_document_versions(base_document_id)
        
        version_info = next((v for v in versions if v["version"] == version), None)
        
        if not version_info:
            raise HTTPException(status_code=404, detail=f"版本 {version} 不存在")
        
        return {
            "base_document_id": base_document_id,
            "version": version_info
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"查询版本详情失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"查询失败: {str(e)}")


@router.get("/{base_document_id}/compare")
async def compare_document_versions(
    base_document_id: str,
    version1: str = Query(..., description="第一个版本号（如 V1）"),
    version2: str = Query(..., description="第二个版本号（如 V3）"),
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    比较两个版本的差异
    
    Args:
        base_document_id: 基础文档ID（group_id）
        version1: 第一个版本号
        version2: 第二个版本号
    
    Returns:
        版本比较结果，包含实体和关系的变化
    """
    try:
        result = VersionComparisonService.compare_versions(
            base_document_id,
            version1,
            version2
        )
        
        return result
    except Exception as e:
        logger.error(f"版本比较失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"比较失败: {str(e)}")


@router.delete("/{document_id}", status_code=200)
async def delete_document(
    document_id: str,
    provider: str = Query("qianwen", description="LLM提供商")
):
    """
    删除文档及其所有相关的Episode、Entity和Relationship
    
    Args:
        document_id: 文档ID（group_id）
        provider: LLM提供商
    
    Returns:
        删除统计信息
    """
    try:
        from app.core.neo4j_client import neo4j_client
        
        # 统计要删除的数据
        stats_query = """
        MATCH (e:Episodic)
        WHERE e.group_id = $document_id
        WITH collect(e.uuid) as episode_uuids
        
        MATCH (n:Entity)
        WHERE n.group_id = $document_id
        
        MATCH (a)-[r:RELATES_TO|MENTIONS]->(b)
        WHERE r.group_id = $document_id OR r.episode_uuid IN episode_uuids
        
        RETURN 
          size(episode_uuids) as episode_count,
          count(DISTINCT n) as entity_count,
          count(DISTINCT r) as relationship_count
        """
        
        stats_result = neo4j_client.execute_query(stats_query, {"document_id": document_id})
        stats = stats_result[0] if stats_result else {}
        episode_count = stats.get("episode_count", 0)
        entity_count = stats.get("entity_count", 0)
        relationship_count = stats.get("relationship_count", 0)
        
        # 删除所有相关的关系（先删除关系，避免约束问题）
        delete_relationships_query = """
        MATCH (e:Episodic)
        WHERE e.group_id = $document_id
        WITH collect(e.uuid) as episode_uuids
        
        MATCH (a)-[r:RELATES_TO|MENTIONS]->(b)
        WHERE r.group_id = $document_id OR r.episode_uuid IN episode_uuids
        DELETE r
        RETURN count(r) as deleted_count
        """
        
        neo4j_client.execute_write(delete_relationships_query, {"document_id": document_id})
        logger.info(f"已删除 {relationship_count} 个关系")
        
        # 删除所有相关的Entity节点
        delete_entities_query = """
        MATCH (n:Entity)
        WHERE n.group_id = $document_id
        DETACH DELETE n
        RETURN count(n) as deleted_count
        """
        
        neo4j_client.execute_write(delete_entities_query, {"document_id": document_id})
        logger.info(f"已删除 {entity_count} 个实体")
        
        # 删除所有相关的Episode节点
        delete_episodes_query = """
        MATCH (e:Episodic)
        WHERE e.group_id = $document_id
        DETACH DELETE e
        RETURN count(e) as deleted_count
        """
        
        neo4j_client.execute_write(delete_episodes_query, {"document_id": document_id})
        logger.info(f"已删除 {episode_count} 个Episode")
        
        return {
            "success": True,
            "document_id": document_id,
            "deleted": {
                "episodes": episode_count,
                "entities": entity_count,
                "relationships": relationship_count
            },
            "message": f"成功删除文档，共删除 {episode_count} 个Episode、{entity_count} 个实体、{relationship_count} 个关系"
        }
    except Exception as e:
        logger.error(f"删除文档失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"删除失败: {str(e)}")

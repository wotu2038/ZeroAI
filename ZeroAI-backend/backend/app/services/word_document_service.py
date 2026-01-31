"""
Word 文档处理服务

实现 Word 文档的解析、分块和 Episode 创建
"""
from docx import Document
from typing import Dict, List, Any, Optional
from datetime import datetime
import os
import logging
import uuid
from app.core.graphiti_client import get_graphiti_instance
from app.models.graphiti_entities import ENTITY_TYPES, EDGE_TYPES, EDGE_TYPE_MAP

logger = logging.getLogger(__name__)


class WordDocumentService:
    """Word 文档处理服务"""
    
    @staticmethod
    def _parse_word_document(file_path: str, document_id: str = None) -> Dict[str, Any]:
        """
        解析 Word 文档，提取文字、图片、链接、表格
        
        Returns:
            {
                "text_content": str,
                "structured_content": list,
                "images": list,
                "links": list,
                "tables": list,
                "metadata": dict
            }
        """
        doc = Document(file_path)
        
        result = {
            "text_content": "",
            "structured_content": [],
            "images": [],
            "links": [],
            "ole_objects": [],  # 新增：OLE对象（嵌入文档）
            "tables": [],
            "metadata": {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified,
            }
        }
        
        # 导入必要的类
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        
        # 先提取所有图片，建立位置映射
        result["images"] = WordDocumentService._extract_images_from_document(doc, document_id, file_path)
        image_position_map = {}  # 记录图片在文档中的位置（段落索引 -> 图片列表）
        unmatched_images = []  # 记录未匹配位置的图片
        
        # 表格计数器（按文档顺序分配table_id）
        table_counter = 0
        # OLE对象计数器（按文档顺序分配ole_id）
        ole_counter = 0
        
        for img in result["images"]:
            pos = img.get("position", -1)
            # 如果position是-1，说明图片未匹配到位置，先记录到unmatched_images
            if pos == -1:
                unmatched_images.append(img)
                logger.warning(f"图片 {img.get('image_id', 'unknown')} 未匹配到段落位置 (rel_id: {img.get('rel_id', 'None')})")
            else:
                # position >= 0 的图片正常映射
                if pos not in image_position_map:
                    image_position_map[pos] = []
                image_position_map[pos].append(img)
        
        # 对于未匹配的图片，尝试通过顺序推断位置（备用策略）
        if unmatched_images:
            logger.info(f"发现 {len(unmatched_images)} 张未匹配位置的图片，将使用备用策略")
            # 这些图片将在后续处理中，根据它们在文档中的出现顺序来推断位置
        
        # 构建章节标题映射（用于图片上下文和描述生成）
        section_titles = []
        for para_idx_temp, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading'):
                section_titles.append((para_idx_temp, para.text.strip()))
        logger.debug(f"构建章节标题映射: {len(section_titles)} 个章节")
        
        # 解析文档内容
        para_idx = 0  # 段落索引计数器
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # 段落
                paragraph = Paragraph(element, doc)
                # 提取格式化的文本（保留加粗、斜体等格式）
                text = WordDocumentService._extract_formatted_text(paragraph)
                
                # 检查是否是标题
                is_heading = paragraph.style.name.startswith('Heading')
                heading_level = 0
                if is_heading:
                    try:
                        heading_level = int(paragraph.style.name.split()[-1])
                    except:
                        heading_level = 1
                
                # 提取链接
                links = WordDocumentService._extract_links_from_paragraph(paragraph)
                if links:
                    result["links"].extend(links)
                
                # 提取OLE对象（嵌入文档）
                ole_objects = WordDocumentService._extract_ole_objects_from_paragraph(paragraph, para_idx, document_id, file_path)
                if ole_objects:
                    # 为每个OLE对象分配固定的ole_id（按文档顺序）
                    for ole_obj in ole_objects:
                        ole_counter += 1
                        ole_id = f"ole_{ole_counter}"
                        ole_obj["ole_id"] = ole_id
                        
                        # 如果有临时文件，需要重命名
                        if ole_obj.get("temp_ole_id") and ole_obj.get("file_path") and ole_obj.get("file_ext"):
                            import os
                            temp_ole_id = ole_obj["temp_ole_id"]
                            old_file_path = ole_obj["file_path"]
                            file_ext = ole_obj["file_ext"]
                            
                            # 构建新文件路径
                            new_file_path = old_file_path.replace(temp_ole_id, ole_id)
                            new_relative_path = ole_obj.get("relative_path", "").replace(temp_ole_id, ole_id)
                            
                            # 重命名文件
                            try:
                                if os.path.exists(old_file_path):
                                    os.rename(old_file_path, new_file_path)
                                    ole_obj["file_path"] = new_file_path
                                    ole_obj["relative_path"] = new_relative_path
                                    logger.debug(f"OLE对象文件重命名: {old_file_path} -> {new_file_path}")
                                    
                                    # 更新显示名称，使用实际保存的扩展名
                                    if file_ext and file_ext != '.bin':
                                        current_name = ole_obj.get("name", "")
                                        # 移除旧的扩展名（.bin或其他）
                                        base_name = os.path.splitext(current_name)[0]
                                        if not base_name or base_name.endswith('.bin') or 'oleObject' in base_name:
                                            base_name = "嵌入文档"
                                        ole_obj["name"] = f"{base_name}{file_ext}"
                                        logger.info(f"更新OLE对象显示名称: {ole_obj['name']} (格式: {file_ext})")
                            except Exception as e:
                                logger.warning(f"重命名OLE对象文件失败: {e}")
                            
                            # 清理临时字段
                            ole_obj.pop("temp_ole_id", None)
                            ole_obj.pop("file_ext", None)
                    
                    result["ole_objects"].extend(ole_objects)
                
                # 检查当前段落位置是否有图片
                images_in_para = image_position_map.get(para_idx, [])
                
                # 对于未匹配的图片，如果当前段落包含图片相关的关键词，尝试关联
                # 这是一个备用策略，用于处理无法通过关系ID匹配的图片
                if unmatched_images and not images_in_para:
                    # 检查段落文本中是否包含图片相关的关键词
                    para_text_lower = text.lower()
                    image_keywords = ['图', '流程图', '示意图', '图片', '图表', 'figure', 'image']
                    if any(keyword in para_text_lower for keyword in image_keywords):
                        # 如果段落包含图片关键词，且还有未匹配的图片，尝试关联第一个未匹配的图片
                        if unmatched_images:
                            img = unmatched_images.pop(0)
                            
                            # 获取前后段落文本（增强版：多段落）
                            prev_paras_text = []
                            next_paras_text = []
                            for i in range(max(0, para_idx - 2), para_idx):
                                if i < len(doc.paragraphs):
                                    prev_text = doc.paragraphs[i].text.strip()
                                    if prev_text:
                                        prev_paras_text.append(prev_text)
                            for i in range(para_idx + 1, min(para_idx + 3, len(doc.paragraphs))):
                                if i < len(doc.paragraphs):
                                    next_text = doc.paragraphs[i].text.strip()
                                    if next_text:
                                        next_paras_text.append(next_text)
                            
                            # 获取最近的章节标题
                            nearest_section_title = ""
                            for section_idx, section_title in reversed(section_titles):
                                if section_idx <= para_idx:
                                    nearest_section_title = section_title
                                    break
                            
                            # 使用增强的描述生成函数
                            description = WordDocumentService._generate_image_description(
                                text, prev_paras_text, next_paras_text, nearest_section_title
                            )
                            
                            # 计算相对位置
                            total_paragraphs = len(doc.paragraphs)
                            relative_position = para_idx / total_paragraphs if total_paragraphs > 0 else 0.0
                            
                            img["position"] = para_idx
                            img["description"] = description
                            img["context"] = text[:300] if text else ""
                            img["prev_context"] = " | ".join(prev_paras_text[:2])[:200] if prev_paras_text else ""
                            img["next_context"] = " | ".join(next_paras_text[:2])[:200] if next_paras_text else ""
                            img["section_title"] = nearest_section_title
                            img["relative_position"] = relative_position
                            img["match_method"] = "keyword"
                            img["match_confidence"] = 0.6  # 关键词匹配的置信度
                            
                            # 添加到当前段落的图片列表
                            images_in_para = [img]
                            logger.info(f"🔍 图片 {img.get('image_id')} 通过关键词匹配到段落 {para_idx}（置信度: 0.6）")
                
                # 合并OLE对象到当前段落
                ole_in_para = [obj for obj in ole_objects]
                
                # 如果有文本，添加到结构化内容
                if text:
                    # 添加到结构化内容
                    result["structured_content"].append({
                        "type": "heading" if is_heading else "paragraph",
                        "level": heading_level if is_heading else 0,
                        "text": text,
                        "links": links,
                        "images": images_in_para,  # 关联图片
                        "ole_objects": ole_in_para  # 关联OLE对象
                    })
                    
                    # 添加到纯文本
                    result["text_content"] += text + "\n"
                    
                    # 如果有图片，添加图片描述到文本
                    if images_in_para:
                        for img in images_in_para:
                            result["text_content"] += f"\n[图片: {img.get('description', '图片')}]\n"
                    
                    # 如果有OLE对象，添加嵌入文档描述到文本
                    if ole_in_para:
                        for ole in ole_in_para:
                            result["text_content"] += f"\n[嵌入文档: {ole.get('name', '文档')} ({ole.get('type', '未知类型')})]\n"
                
                # 如果段落没有文本但有图片或OLE对象，单独记录
                elif images_in_para or ole_in_para:
                    result["structured_content"].append({
                        "type": "image_only" if images_in_para else "ole_only",
                        "level": 0,
                        "text": "",
                        "links": links,
                        "images": images_in_para,
                        "ole_objects": ole_in_para
                    })
                    for img in images_in_para:
                        result["text_content"] += f"\n[图片: {img.get('description', '图片')}]\n"
                    for ole in ole_in_para:
                        result["text_content"] += f"\n[嵌入文档: {ole.get('name', '文档')} ({ole.get('type', '未知类型')})]\n"
                
                para_idx += 1  # 增加段落索引
            
            elif isinstance(element, CT_Tbl):
                # 表格 - 按文档顺序分配table_id
                table = Table(element, doc)
                table_data = WordDocumentService._extract_table_data(table)
                
                # 分配固定的table_id（按文档顺序）
                table_counter += 1
                table_id = f"table_{table_counter}"
                table_data["table_id"] = table_id
                
                result["tables"].append(table_data)
                
                # 将表格转换为文本描述
                table_text = WordDocumentService._format_table_as_text(table_data)
                result["text_content"] += table_text + "\n"
                result["structured_content"].append({
                    "type": "table",
                    "data": table_data,
                    "text": table_text,
                    "table_id": table_id  # 保存table_id
                })
        
        # 处理仍未匹配的图片（最后的备用策略）
        if unmatched_images:
            logger.warning(f"仍有 {len(unmatched_images)} 张图片未匹配到位置，将关联到文档末尾")
            # 将这些图片关联到最后一个段落（作为最后的备用策略）
            last_para_idx = para_idx - 1 if para_idx > 0 else 0
            for img in unmatched_images:
                img["position"] = last_para_idx
                logger.warning(f"图片 {img.get('image_id')} 未匹配到位置，已关联到段落 {last_para_idx}（备用策略）")
                # 确保这些图片也被添加到image_position_map中
                if last_para_idx not in image_position_map:
                    image_position_map[last_para_idx] = []
                image_position_map[last_para_idx].append(img)
                # 同时添加到最后一个structured_content项中（如果存在）
                if result["structured_content"]:
                    last_item = result["structured_content"][-1]
                    if "images" not in last_item:
                        last_item["images"] = []
                    last_item["images"].append(img)
                    # 也添加到text_content
                    result["text_content"] += f"\n[图片: {img.get('description', '图片')}]\n"
        
        return result
    
    @staticmethod
    def _extract_ole_objects_from_paragraph(paragraph, para_idx: int, document_id: str = None, file_path: str = None) -> List[Dict]:
        """从段落中提取OLE对象（嵌入文档），支持Microsoft Office和WPS格式"""
        ole_objects = []
        seen_r_ids = set()  # 用于去重，避免同一个r_id被多次提取
        try:
            from docx.oxml.ns import qn
            import xml.etree.ElementTree as ET
            
            # 遍历段落中的所有runs
            for run in paragraph.runs:
                # 方法1：检查是否有OLEObject元素（Microsoft Office格式）
                # OLEObject 在 o: 命名空间中
                try:
                    ole_elements = run._element.xpath('.//o:OLEObject', namespaces={
                        'o': 'urn:schemas-microsoft-com:office:office'
                    })
                except:
                    # 如果命名空间失败，尝试使用local-name
                    ole_elements = run._element.xpath('.//*[local-name()="OLEObject"]')
                
                for ole_element in ole_elements:
                    # 获取OLE对象信息
                    prog_id = ole_element.get('ProgId', '')
                    ole_type = ole_element.get('Type', '')
                    r_id = ole_element.get(qn('r:id'))
                    
                    # 去重：如果这个r_id已经被处理过，跳过
                    if r_id and r_id in seen_r_ids:
                        logger.debug(f"跳过重复的OLE对象: r_id={r_id}, 段落={para_idx}")
                        continue
                    
                    if r_id:
                        seen_r_ids.add(r_id)
                    
                    # 从关系ID获取嵌入文档信息
                    file_name = ""
                    file_type = "未知类型"
                    
                    if r_id:
                        try:
                            rel = paragraph.part.rels[r_id]
                            if hasattr(rel, 'target_ref'):
                                file_name = rel.target_ref
                            else:
                                file_name = str(rel.target)
                            
                            # 根据ProgId判断文件类型
                            if 'Excel' in prog_id or 'excel' in prog_id.lower():
                                file_type = "Excel文件"
                            elif 'Word' in prog_id or 'word' in prog_id.lower():
                                file_type = "Word文档"
                            elif 'PowerPoint' in prog_id or 'powerpoint' in prog_id.lower():
                                file_type = "PowerPoint演示文稿"
                            elif 'PDF' in prog_id or 'pdf' in prog_id.lower():
                                file_type = "PDF文档"
                            else:
                                file_type = prog_id or "嵌入对象"
                            
                            # 如果没有文件名，使用ProgId
                            if not file_name:
                                file_name = prog_id or "嵌入文档"
                        except (KeyError, AttributeError) as e:
                            logger.debug(f"提取OLE对象关系失败: {e}, r_id={r_id}")
                            file_name = prog_id or "嵌入文档"
                    
                    # 尝试提取并保存嵌入文档（与WPS格式相同的逻辑）
                    saved_file_path = None
                    relative_path = None
                    actual_ext = '.bin'
                    temp_ole_id = None
                    
                    if r_id and file_path and document_id:
                        try:
                            rel = paragraph.part.rels[r_id]
                            # 从docx的zip文件中提取嵌入文档
                            import zipfile
                            
                            # 获取嵌入文档路径
                            embed_path = None
                            if hasattr(rel, 'target_ref'):
                                embed_path = rel.target_ref
                            elif hasattr(rel, 'target'):
                                embed_path = str(rel.target)
                            
                            if embed_path:
                                with zipfile.ZipFile(file_path, 'r') as zip_file:
                                    # 尝试不同的路径格式
                                    possible_paths = [
                                        embed_path,
                                        f"word/{embed_path}",
                                        f"word/embeddings/{os.path.basename(embed_path)}",
                                        embed_path.replace('../', ''),
                                        embed_path.replace('embeddings/', 'word/embeddings/'),
                                    ]
                                    
                                    found_path = None
                                    for path in possible_paths:
                                        if path in zip_file.namelist():
                                            found_path = path
                                            break
                                    
                                    if found_path:
                                        # 创建保存目录
                                        ole_dir = os.path.abspath(f"uploads/extracted_ole/{document_id}")
                                        os.makedirs(ole_dir, exist_ok=True)
                                        
                                        # 先读取文件内容，检测实际格式
                                        with zip_file.open(found_path) as source:
                                            file_content = source.read()
                                        
                                        # 检测文件格式（可能会提取并保存文件）
                                        temp_ole_id = f"ole_temp_{para_idx}_{len(ole_objects)}"
                                        actual_ext = WordDocumentService._detect_file_format(file_content, found_path, ole_dir, temp_ole_id, prog_id)
                                        
                                        # 检查是否已经保存了提取的文件（标准格式）
                                        extracted_file = os.path.join(ole_dir, f"{temp_ole_id}{actual_ext}")
                                        
                                        if os.path.exists(extracted_file) and actual_ext in ['.xls', '.xlsx', '.doc', '.docx', '.ppt', '.pptx']:
                                            # 文件已经由_detect_file_format提取并保存为标准格式
                                            saved_file_path = extracted_file
                                            relative_path = f"extracted_ole/{document_id}/{temp_ole_id}{actual_ext}"
                                            logger.info(f"✓ 使用已提取的标准格式文件: {saved_file_path} (格式: {actual_ext})")
                                        elif actual_ext in ['.xls', '.xlsx', '.doc', '.docx', '.ppt', '.pptx']:
                                            # _detect_file_format返回了标准格式，但文件不存在（可能是提取失败）
                                            logger.warning(f"检测到格式为{actual_ext}，但提取的文件不存在，尝试强制提取...")
                                            retry_temp_id = f"{temp_ole_id}_retry"
                                            retry_ext = WordDocumentService._detect_file_format(file_content, found_path, ole_dir, retry_temp_id, prog_id)
                                            retry_file = os.path.join(ole_dir, f"{retry_temp_id}{retry_ext}")
                                            
                                            if os.path.exists(retry_file) and retry_ext in ['.xls', '.xlsx', '.doc', '.docx', '.ppt', '.pptx']:
                                                saved_file_path = os.path.join(ole_dir, f"{temp_ole_id}{retry_ext}")
                                                os.rename(retry_file, saved_file_path)
                                                relative_path = f"extracted_ole/{document_id}/{temp_ole_id}{retry_ext}"
                                                actual_ext = retry_ext
                                                logger.info(f"✓ 强制提取成功: {saved_file_path} (格式: {retry_ext})")
                                            else:
                                                # 提取失败，保存为.bin格式
                                                actual_ext = '.bin'
                                                saved_file_path = os.path.join(ole_dir, f"{temp_ole_id}{actual_ext}")
                                                relative_path = f"extracted_ole/{document_id}/{temp_ole_id}{actual_ext}"
                                                with open(saved_file_path, 'wb') as target:
                                                    target.write(file_content)
                                                logger.warning(f"✗ 提取失败，保存为原始.bin格式: {saved_file_path}")
                                        else:
                                            # 返回的是.bin或其他格式，直接保存原始文件
                                            saved_file_path = os.path.join(ole_dir, f"{temp_ole_id}{actual_ext}")
                                            relative_path = f"extracted_ole/{document_id}/{temp_ole_id}{actual_ext}"
                                            
                                            with open(saved_file_path, 'wb') as target:
                                                target.write(file_content)
                                            
                                            logger.info(f"保存嵌入文档: {saved_file_path} (来源: {found_path}, 格式: {actual_ext})")
                                    else:
                                        logger.debug(f"在zip文件中未找到嵌入文档路径: {embed_path}, 尝试的路径: {possible_paths}")
                        except Exception as e:
                            logger.warning(f"提取嵌入文档文件失败: {e}", exc_info=True)
                    
                    # 获取上下文（段落文本）
                    para_text = paragraph.text.strip()
                    
                    # 更新文件名，使用实际保存的扩展名
                    if actual_ext and actual_ext != '.bin':
                        base_name = os.path.splitext(file_name)[0] if file_name else "嵌入文档"
                        display_name = f"{base_name}{actual_ext}"
                    else:
                        display_name = file_name
                    
                    # ole_id将在调用处按文档顺序分配，这里先不设置
                    ole_objects.append({
                        "position": para_idx,
                        "name": display_name,  # 使用更新后的文件名
                        "type": file_type,
                        "prog_id": prog_id,
                        "ole_type": ole_type,
                        "context": para_text[:100] if para_text else "",
                        "file_path": saved_file_path,
                        "relative_path": relative_path,
                        "temp_ole_id": temp_ole_id,  # 保存临时ID，用于后续重命名文件
                        "file_ext": actual_ext  # 保存文件扩展名
                    })
                
                # 方法2：检查是否有Object元素（另一种嵌入方式）
                try:
                    object_elements = run._element.xpath('.//o:Object', namespaces={
                        'o': 'urn:schemas-microsoft-com:office:office'
                    })
                except:
                    object_elements = run._element.xpath('.//*[local-name()="Object"]')
                
                # 方法3：检查WPS格式的嵌入对象（通过XML内容检查）
                try:
                    run_xml = run._element.xml
                    if run_xml:
                        # WPS可能使用不同的标签或命名空间
                        # 检查是否包含Excel相关的关键词
                        if any(keyword in run_xml.lower() for keyword in ['excel', 'xls', 'xlsx', 'spreadsheet']):
                            # 尝试从XML中提取信息
                            root = ET.fromstring(run_xml)
                            for elem in root.iter():
                                tag = elem.tag.lower()
                                if 'excel' in tag or 'ole' in tag or 'object' in tag or 'embed' in tag:
                                    # 提取ProgId或类型信息
                                    prog_id = elem.get('ProgId', '') or elem.get('progid', '')
                                    if not prog_id:
                                        # 尝试从属性中获取
                                        for attr_name, attr_value in elem.attrib.items():
                                            if 'excel' in attr_value.lower() or 'xls' in attr_value.lower():
                                                prog_id = attr_value
                                                break
                                    
                                    if prog_id or 'excel' in tag or 'xls' in tag:
                                        # 获取关系ID
                                        r_id = None
                                        for attr_name, attr_value in elem.attrib.items():
                                            if 'id' in attr_name.lower() and attr_value:
                                                try:
                                                    r_id = elem.get(qn('r:id')) or attr_value
                                                    break
                                                except:
                                                    r_id = attr_value
                                                    break
                                        
                                        # 去重：如果这个r_id已经被处理过，跳过
                                        if r_id and r_id in seen_r_ids:
                                            logger.debug(f"跳过重复的WPS格式OLE对象: r_id={r_id}, 段落={para_idx}")
                                            continue
                                        
                                        if r_id:
                                            seen_r_ids.add(r_id)
                                        
                                        file_name = ""
                                        file_type = "Excel文件"
                                        
                                        if r_id:
                                            try:
                                                rel = paragraph.part.rels[r_id]
                                                if hasattr(rel, 'target_ref'):
                                                    file_name = rel.target_ref
                                                else:
                                                    file_name = str(rel.target) if hasattr(rel, 'target') else ""
                                            except:
                                                pass
                                        
                                        if not file_name:
                                            file_name = prog_id or "Excel文件"
                                        
                                        # 尝试提取并保存嵌入文档
                                        saved_file_path = None
                                        relative_path = None
                                        if r_id and file_path and document_id:
                                            try:
                                                rel = paragraph.part.rels[r_id]
                                                # 从docx的zip文件中提取嵌入文档
                                                import zipfile
                                                import shutil
                                                
                                                # 获取嵌入文档路径
                                                embed_path = None
                                                if hasattr(rel, 'target_ref'):
                                                    embed_path = rel.target_ref
                                                elif hasattr(rel, 'target'):
                                                    embed_path = str(rel.target)
                                                
                                                if embed_path:
                                                    with zipfile.ZipFile(file_path, 'r') as zip_file:
                                                        # 尝试不同的路径格式
                                                        possible_paths = [
                                                            embed_path,
                                                            f"word/{embed_path}",
                                                            f"word/embeddings/{os.path.basename(embed_path)}",
                                                            embed_path.replace('../', ''),
                                                            embed_path.replace('embeddings/', 'word/embeddings/'),
                                                        ]
                                                        
                                                        found_path = None
                                                        for path in possible_paths:
                                                            if path in zip_file.namelist():
                                                                found_path = path
                                                                break
                                                        
                                                        if found_path:
                                                            # 创建保存目录
                                                            ole_dir = os.path.abspath(f"uploads/extracted_ole/{document_id}")
                                                            os.makedirs(ole_dir, exist_ok=True)
                                                            
                                                            # 先读取文件内容，检测实际格式
                                                            with zip_file.open(found_path) as source:
                                                                file_content = source.read()
                                                            
                                                            # 检测文件格式（可能会提取并保存文件）
                                                            # ole_id将在调用处按文档顺序分配，这里使用临时ID用于文件保存
                                                            temp_ole_id = f"ole_temp_{para_idx}_{len(ole_objects)}"
                                                            actual_ext = WordDocumentService._detect_file_format(file_content, found_path, ole_dir, temp_ole_id, prog_id)
                                                            
                                                            # 检查是否已经保存了提取的文件（标准格式）
                                                            extracted_file = os.path.join(ole_dir, f"{temp_ole_id}{actual_ext}")
                                                            
                                                            if os.path.exists(extracted_file) and actual_ext in ['.xls', '.xlsx', '.doc', '.docx', '.ppt', '.pptx']:
                                                                # 文件已经由_detect_file_format提取并保存为标准格式
                                                                saved_file_path = extracted_file
                                                                relative_path = f"extracted_ole/{document_id}/{temp_ole_id}{actual_ext}"
                                                                logger.info(f"✓ 使用已提取的标准格式文件: {saved_file_path} (格式: {actual_ext})")
                                                            elif actual_ext in ['.xls', '.xlsx', '.doc', '.docx', '.ppt', '.pptx']:
                                                                # _detect_file_format返回了标准格式，但文件不存在（可能是提取失败）
                                                                # 尝试强制提取：重新调用_detect_file_format，确保提取成功
                                                                logger.warning(f"检测到格式为{actual_ext}，但提取的文件不存在，尝试强制提取...")
                                                                # 重新尝试提取（使用新的临时ID避免冲突）
                                                                retry_temp_id = f"{temp_ole_id}_retry"
                                                                retry_ext = WordDocumentService._detect_file_format(file_content, found_path, ole_dir, retry_temp_id, prog_id)
                                                                retry_file = os.path.join(ole_dir, f"{retry_temp_id}{retry_ext}")
                                                                
                                                                if os.path.exists(retry_file) and retry_ext in ['.xls', '.xlsx', '.doc', '.docx', '.ppt', '.pptx']:
                                                                    # 重命名为原始ID
                                                                    saved_file_path = os.path.join(ole_dir, f"{temp_ole_id}{retry_ext}")
                                                                    os.rename(retry_file, saved_file_path)
                                                                    relative_path = f"extracted_ole/{document_id}/{temp_ole_id}{retry_ext}"
                                                                    logger.info(f"✓ 强制提取成功: {saved_file_path} (格式: {retry_ext})")
                                                                else:
                                                                    # 提取失败，保存为.bin格式
                                                                    actual_ext = '.bin'
                                                                    saved_file_path = os.path.join(ole_dir, f"{temp_ole_id}{actual_ext}")
                                                                    relative_path = f"extracted_ole/{document_id}/{temp_ole_id}{actual_ext}"
                                                                    with open(saved_file_path, 'wb') as target:
                                                                        target.write(file_content)
                                                                    logger.warning(f"✗ 提取失败，保存为原始.bin格式: {saved_file_path}")
                                                            else:
                                                                # 返回的是.bin或其他格式，直接保存原始文件
                                                                saved_file_path = os.path.join(ole_dir, f"{temp_ole_id}{actual_ext}")
                                                                relative_path = f"extracted_ole/{document_id}/{temp_ole_id}{actual_ext}"
                                                                
                                                                with open(saved_file_path, 'wb') as target:
                                                                    target.write(file_content)
                                                                
                                                                logger.info(f"保存嵌入文档: {saved_file_path} (来源: {found_path}, 格式: {actual_ext})")
                                                        else:
                                                            logger.debug(f"在zip文件中未找到嵌入文档路径: {embed_path}, 尝试的路径: {possible_paths}")
                                            except Exception as e:
                                                logger.warning(f"提取嵌入文档文件失败: {e}", exc_info=True)
                                        
                                        para_text = paragraph.text.strip()
                                        
                                        # ole_id将在调用处按文档顺序分配，这里先不设置
                                        # 更新文件名，使用实际保存的扩展名
                                        if actual_ext and actual_ext != '.bin':
                                            # 如果成功提取为标准格式，更新文件名显示
                                            base_name = os.path.splitext(file_name)[0] if file_name else "嵌入文档"
                                            display_name = f"{base_name}{actual_ext}"
                                        else:
                                            display_name = file_name
                                        
                                        ole_objects.append({
                                            "position": para_idx,
                                            "name": display_name,  # 使用更新后的文件名
                                            "type": file_type,
                                            "prog_id": prog_id or "Excel.Sheet",
                                            "ole_type": "WPS_Embedded",
                                            "context": para_text[:100] if para_text else "",
                                            "file_path": saved_file_path,
                                            "relative_path": relative_path,
                                            "temp_ole_id": temp_ole_id,  # 保存临时ID，用于后续重命名文件
                                            "file_ext": actual_ext  # 保存文件扩展名
                                        })
                                        logger.info(f"从段落 {para_idx} 提取到WPS格式的Excel嵌入对象: {file_name}")
                except Exception as e:
                    logger.debug(f"检查WPS格式嵌入对象时出错: {e}")
                
                for obj_element in object_elements:
                    prog_id = obj_element.get('ProgId', '')
                    r_id = obj_element.get(qn('r:id'))
                    
                    # 去重：如果这个r_id已经被处理过，跳过
                    if r_id and r_id in seen_r_ids:
                        logger.debug(f"跳过重复的Object元素OLE对象: r_id={r_id}, 段落={para_idx}")
                        continue
                    
                    if r_id:
                        seen_r_ids.add(r_id)
                    
                    if prog_id or r_id:
                        file_name = ""
                        file_type = "未知类型"
                        
                        if r_id:
                            try:
                                rel = paragraph.part.rels[r_id]
                                if hasattr(rel, 'target_ref'):
                                    file_name = rel.target_ref
                                else:
                                    file_name = str(rel.target)
                            except (KeyError, AttributeError):
                                file_name = prog_id or "嵌入文档"
                        
                        if 'Excel' in prog_id or 'excel' in prog_id.lower():
                            file_type = "Excel文件"
                        elif 'Word' in prog_id or 'word' in prog_id.lower():
                            file_type = "Word文档"
                        else:
                            file_type = prog_id or "嵌入对象"
                        
                        para_text = paragraph.text.strip()
                        
                        # ole_id将在调用处按文档顺序分配，这里先不设置
                        ole_objects.append({
                            "position": para_idx,
                            "name": file_name or prog_id or "嵌入文档",
                            "type": file_type,
                            "prog_id": prog_id,
                            "ole_type": "Object",
                            "context": para_text[:100] if para_text else ""
                        })
            
            if ole_objects:
                logger.info(f"从段落 {para_idx} 提取到 {len(ole_objects)} 个OLE对象")
        except Exception as e:
            logger.warning(f"从段落提取OLE对象时出错: {e}", exc_info=True)
        
        return ole_objects
    
    @staticmethod
    def _extract_links_from_paragraph(paragraph) -> List[Dict]:
        """从段落中提取链接"""
        links = []
        try:
            from docx.oxml.ns import qn
            
            # 遍历段落中的所有runs
            for run in paragraph.runs:
                # 检查是否有超链接
                hyperlinks = run._element.xpath('.//w:hyperlink')
                
                for hyperlink in hyperlinks:
                    # 获取链接地址
                    r_id = hyperlink.get(qn('r:id'))
                    if r_id:
                        # 从文档的关系中获取链接地址
                        try:
                            rel = paragraph.part.rels[r_id]
                            url = rel.target_ref if hasattr(rel, 'target_ref') else str(rel.target)
                            
                            # 获取链接文本
                            link_text = run.text.strip()
                            if not link_text:
                                # 如果没有文本，尝试从hyperlink元素中获取
                                link_text = ''.join(hyperlink.itertext()).strip()
                            
                            if url:
                                # 判断链接类型
                                link_type = "external"
                                if url.startswith('#'):
                                    link_type = "internal"
                                elif url.startswith('file://') or url.endswith(('.docx', '.doc', '.xlsx', '.xls', '.pdf')):
                                    link_type = "file"
                                
                                links.append({
                                    "text": link_text or url,  # 如果没有文本，使用URL
                                    "url": url,
                                    "type": link_type
                                })
                        except (KeyError, AttributeError) as e:
                            logger.debug(f"提取链接失败: {e}, r_id={r_id}")
                            continue
        except Exception as e:
            logger.warning(f"从段落提取链接时出错: {e}", exc_info=True)
        
        return links
    
    @staticmethod
    def _extract_table_data(table) -> Dict[str, Any]:
        """提取表格数据"""
        if not table.rows:
            return {"headers": [], "rows": [], "row_count": 0, "col_count": 0}
        
        # 提取表头（第一行）
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # 提取数据行
        rows = []
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)
        
        return {
            "headers": headers,
            "rows": rows,
            "row_count": len(rows),
            "col_count": len(headers)
        }
    
    @staticmethod
    def _detect_file_format(file_content: bytes, original_path: str, save_dir: str, ole_id: str, prog_id: str = None) -> str:
        """
        检测文件的实际格式并提取内容
        
        Args:
            file_content: 文件内容（字节）
            original_path: 原始文件路径
            save_dir: 保存目录
            ole_id: OLE对象ID
            prog_id: OLE对象的ProgId（用于辅助判断文件类型）
            
        Returns:
            实际的文件扩展名（如 .xlsx, .xls, .docx, .doc, .pptx, .ppt等）
        """
        if len(file_content) < 8:
            return '.bin'
        
        header = file_content[:8]
        
        # 检查是否是ZIP格式（.xlsx, .docx, .pptx实际上是ZIP）
        if header[:2] == b'PK':
            # 尝试作为ZIP打开，检查文件类型
            try:
                import zipfile
                import io
                with zipfile.ZipFile(io.BytesIO(file_content), 'r') as zf:
                    file_list = zf.namelist()
                    # 检查是否包含Excel的特征文件
                    if any('xl/' in name or 'xl/workbook' in name or 'xl/worksheets' in name for name in file_list):
                        logger.info(f"检测到Excel格式（.xlsx）: {ole_id}")
                        return '.xlsx'
                    # 检查是否包含Word的特征文件
                    elif any('word/' in name or 'word/document' in name for name in file_list):
                        logger.info(f"检测到Word格式（.docx）: {ole_id}")
                        return '.docx'
                    # 检查是否包含PowerPoint的特征文件
                    elif any('ppt/' in name or 'ppt/presentation' in name or 'ppt/slides' in name for name in file_list):
                        logger.info(f"检测到PowerPoint格式（.pptx）: {ole_id}")
                        return '.pptx'
            except Exception as e:
                logger.debug(f"ZIP格式检测失败: {e}")
            # 如果是ZIP但不是已知的Office格式，返回.zip
            return '.zip'
        
        # 检查是否是OLE2格式（Composite Document File V2）
        # OLE2签名: D0 CF 11 E0 A1 B1 1A E1
        ole2_signature = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'
        if header == ole2_signature:
            # 这是OLE2格式，尝试使用olefile库提取内容
            try:
                import olefile
                import io
                if olefile.isOleFile(io.BytesIO(file_content)):
                    ole = olefile.OleFileIO(io.BytesIO(file_content))
                    try:
                        # 获取所有流
                        stream_list = ole.listdir()
                        
                        # 根据ProgId和流名判断文件类型
                        detected_type = None
                        extracted_data = None
                        stream_name_used = None
                        
                        # 记录所有流名用于调试
                        all_streams_debug = []
                        for s in stream_list:
                            if isinstance(s, tuple):
                                all_streams_debug.append('/'.join(s))
                            else:
                                all_streams_debug.append(str(s))
                        logger.info(f"OLE2文件中的所有流: {all_streams_debug}")
                        
                        # 检查Excel流
                        for stream_name in stream_list:
                            if isinstance(stream_name, tuple) and len(stream_name) > 0:
                                stream_name_str = stream_name[0]
                            else:
                                stream_name_str = str(stream_name)
                            
                            if stream_name_str in ['Workbook', 'Book']:
                                detected_type = '.xls'
                                stream_name_used = stream_name_str
                                break
                        
                        # 检查Word流
                        if not detected_type:
                            for stream_name in stream_list:
                                if isinstance(stream_name, tuple) and len(stream_name) > 0:
                                    stream_name_str = stream_name[0]
                                else:
                                    stream_name_str = str(stream_name)
                                
                                if stream_name_str == 'WordDocument':
                                    detected_type = '.doc'
                                    stream_name_used = stream_name_str
                                    break
                        
                        # 检查PowerPoint流
                        if not detected_type:
                            for stream_name in stream_list:
                                if isinstance(stream_name, tuple) and len(stream_name) > 0:
                                    stream_name_str = stream_name[0]
                                else:
                                    stream_name_str = str(stream_name)
                                
                                if 'PowerPoint' in stream_name_str or 'Presentation' in stream_name_str:
                                    detected_type = '.ppt'
                                    stream_name_used = stream_name_str
                                    break
                        
                        # 如果通过流名无法判断，尝试根据ProgId判断
                        if not detected_type and prog_id:
                            prog_id_lower = prog_id.lower()
                            if 'excel' in prog_id_lower:
                                detected_type = '.xls'
                            elif 'word' in prog_id_lower:
                                detected_type = '.doc'
                            elif 'powerpoint' in prog_id_lower or 'ppt' in prog_id_lower:
                                detected_type = '.ppt'
                        
                        # 如果找到了类型，提取内容
                        if detected_type and stream_name_used:
                            try:
                                # 找到对应的流对象（stream_name_used是字符串，需要找到实际的流对象）
                                actual_stream_name = None
                                for stream_name in stream_list:
                                    if isinstance(stream_name, tuple) and len(stream_name) > 0:
                                        stream_name_str = stream_name[0]
                                    else:
                                        stream_name_str = str(stream_name)
                                    
                                    if stream_name_str == stream_name_used:
                                        actual_stream_name = stream_name
                                        break
                                
                                if actual_stream_name:
                                    extracted_data = ole.openstream(actual_stream_name).read()
                                    # 保存提取的内容
                                    extracted_path = os.path.join(save_dir, f"{ole_id}{detected_type}")
                                    with open(extracted_path, 'wb') as f:
                                        f.write(extracted_data)
                                    logger.info(f"✓ 成功从OLE2格式提取{detected_type}内容: {extracted_path}, 流名: {stream_name_used}")
                                    ole.close()
                                    return detected_type
                                else:
                                    logger.warning(f"✗ 找不到流对象: {stream_name_used}, 可用流: {all_streams_debug}")
                            except Exception as e:
                                logger.warning(f"✗ 提取{detected_type}内容失败: {e}", exc_info=True)
                        
                        # 如果找不到标准流，检查是否有package流（打包的OLE对象）
                        # 即使detected_type存在，如果没有stream_name_used，也要检查package流
                        if not stream_name_used:
                            package_stream = None
                            logger.info(f"开始查找package流，流列表类型: {type(stream_list)}, 数量: {len(stream_list)}")
                            for idx, stream_name in enumerate(stream_list):
                                # 处理流名（可能是tuple或字符串）
                                stream_first_str = None
                                stream_type_info = f"类型: {type(stream_name)}"
                                
                                if isinstance(stream_name, tuple):
                                    stream_type_info += f", 长度: {len(stream_name)}"
                                    if len(stream_name) > 0:
                                        stream_first = stream_name[0]
                                        stream_type_info += f", 第一个元素类型: {type(stream_first)}"
                                        # 处理不同类型的流名
                                        if isinstance(stream_first, bytes):
                                            try:
                                                stream_first_str = stream_first.decode('utf-8', errors='ignore').strip()
                                            except:
                                                stream_first_str = str(stream_first).strip()
                                        elif isinstance(stream_first, str):
                                            stream_first_str = stream_first.strip()
                                        else:
                                            stream_first_str = str(stream_first).strip()
                                else:
                                    if isinstance(stream_name, bytes):
                                        try:
                                            stream_first_str = stream_name.decode('utf-8', errors='ignore').strip()
                                        except:
                                            stream_first_str = str(stream_name).strip()
                                    else:
                                        stream_first_str = str(stream_name).strip()
                                
                                logger.info(f"流[{idx}]: {stream_name} -> {stream_first_str} ({stream_type_info})")
                                
                                # 检查是否是package流（不区分大小写，去除空白字符）
                                # 也检查流名的字符串表示中是否包含'package'
                                if stream_first_str:
                                    stream_lower = stream_first_str.lower()
                                    stream_repr = repr(stream_name).lower()
                                    if stream_lower == 'package' or 'package' in stream_lower or 'package' in stream_repr:
                                        package_stream = stream_name
                                        logger.info(f"✓✓✓ 找到package流: {package_stream} (原始: {stream_first_str}, repr: {stream_repr})")
                                        break
                            
                            if package_stream:
                                logger.info(f"找到package流，尝试提取内容: {package_stream}")
                                try:
                                    package_data = ole.openstream(package_stream).read()
                                    
                                    # 检查提取的数据格式
                                    if len(package_data) >= 2 and package_data[:2] == b'PK':
                                        # ZIP格式，可能是.xlsx、.docx或.pptx文件，需要进一步判断
                                        try:
                                            import zipfile
                                            import io
                                            with zipfile.ZipFile(io.BytesIO(package_data), 'r') as zf:
                                                file_list = zf.namelist()
                                                # 检查是否包含Excel的特征文件
                                                if any('xl/' in name or 'xl/workbook' in name or 'xl/worksheets' in name for name in file_list):
                                                    extracted_path = os.path.join(save_dir, f"{ole_id}.xlsx")
                                                    with open(extracted_path, 'wb') as f:
                                                        f.write(package_data)
                                                    logger.info(f"从package流提取到.xlsx格式文件: {extracted_path}")
                                                    ole.close()
                                                    return '.xlsx'
                                                # 检查是否包含Word的特征文件
                                                elif any('word/' in name or 'word/document' in name for name in file_list):
                                                    extracted_path = os.path.join(save_dir, f"{ole_id}.docx")
                                                    with open(extracted_path, 'wb') as f:
                                                        f.write(package_data)
                                                    logger.info(f"从package流提取到.docx格式文件: {extracted_path}")
                                                    ole.close()
                                                    return '.docx'
                                                # 检查是否包含PowerPoint的特征文件
                                                elif any('ppt/' in name or 'ppt/presentation' in name or 'ppt/slides' in name for name in file_list):
                                                    extracted_path = os.path.join(save_dir, f"{ole_id}.pptx")
                                                    with open(extracted_path, 'wb') as f:
                                                        f.write(package_data)
                                                    logger.info(f"从package流提取到.pptx格式文件: {extracted_path}")
                                                    ole.close()
                                                    return '.pptx'
                                                else:
                                                    # 无法识别，根据ProgId判断
                                                    if prog_id:
                                                        prog_id_lower = prog_id.lower()
                                                        if 'excel' in prog_id_lower:
                                                            extracted_path = os.path.join(save_dir, f"{ole_id}.xlsx")
                                                            with open(extracted_path, 'wb') as f:
                                                                f.write(package_data)
                                                            logger.info(f"从package流提取数据，根据ProgId判断为.xlsx: {extracted_path}")
                                                            ole.close()
                                                            return '.xlsx'
                                                        elif 'word' in prog_id_lower:
                                                            extracted_path = os.path.join(save_dir, f"{ole_id}.docx")
                                                            with open(extracted_path, 'wb') as f:
                                                                f.write(package_data)
                                                            logger.info(f"从package流提取数据，根据ProgId判断为.docx: {extracted_path}")
                                                            ole.close()
                                                            return '.docx'
                                                        elif 'powerpoint' in prog_id_lower or 'ppt' in prog_id_lower:
                                                            extracted_path = os.path.join(save_dir, f"{ole_id}.pptx")
                                                            with open(extracted_path, 'wb') as f:
                                                                f.write(package_data)
                                                            logger.info(f"从package流提取数据，根据ProgId判断为.pptx: {extracted_path}")
                                                            ole.close()
                                                            return '.pptx'
                                                    # 如果无法判断，默认保存为.xlsx（向后兼容）
                                                    extracted_path = os.path.join(save_dir, f"{ole_id}.xlsx")
                                                    with open(extracted_path, 'wb') as f:
                                                        f.write(package_data)
                                                    logger.warning(f"从package流提取ZIP数据，无法识别格式，默认保存为.xlsx: {extracted_path}")
                                                    ole.close()
                                                    return '.xlsx'
                                        except Exception as e:
                                            logger.warning(f"检查ZIP文件内容失败: {e}，根据ProgId判断")
                                            # 如果检查失败，根据ProgId判断
                                            if prog_id:
                                                prog_id_lower = prog_id.lower()
                                                if 'word' in prog_id_lower:
                                                    extracted_path = os.path.join(save_dir, f"{ole_id}.docx")
                                                    with open(extracted_path, 'wb') as f:
                                                        f.write(package_data)
                                                    logger.info(f"从package流提取数据，根据ProgId判断为.docx: {extracted_path}")
                                                    ole.close()
                                                    return '.docx'
                                            # 默认保存为.xlsx
                                            extracted_path = os.path.join(save_dir, f"{ole_id}.xlsx")
                                            with open(extracted_path, 'wb') as f:
                                                f.write(package_data)
                                            logger.warning(f"从package流提取ZIP数据，检查失败，默认保存为.xlsx: {extracted_path}")
                                            ole.close()
                                            return '.xlsx'
                                    elif len(package_data) >= 8 and package_data[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
                                        # OLE2格式，可能是.xls文件，尝试提取Workbook流
                                        try:
                                            import io
                                            package_ole = olefile.OleFileIO(io.BytesIO(package_data))
                                            package_streams = package_ole.listdir()
                                            workbook_found = False
                                            
                                            for pkg_stream in package_streams:
                                                if isinstance(pkg_stream, tuple):
                                                    pkg_stream_first = pkg_stream[0] if len(pkg_stream) > 0 else ''
                                                else:
                                                    pkg_stream_first = str(pkg_stream)
                                                
                                                if pkg_stream_first.lower() in ['workbook', 'book']:
                                                    workbook_data = package_ole.openstream(pkg_stream).read()
                                                    extracted_path = os.path.join(save_dir, f"{ole_id}.xls")
                                                    with open(extracted_path, 'wb') as f:
                                                        f.write(workbook_data)
                                                    logger.info(f"从package流的OLE2格式中提取Workbook流: {extracted_path}")
                                                    package_ole.close()
                                                    ole.close()
                                                    workbook_found = True
                                                    return '.xls'
                                            
                                            if not workbook_found:
                                                # 如果找不到Workbook流，直接保存package数据为.xls
                                                extracted_path = os.path.join(save_dir, f"{ole_id}.xls")
                                                with open(extracted_path, 'wb') as f:
                                                    f.write(package_data)
                                                logger.info(f"从package流提取OLE2数据，保存为.xls: {extracted_path}")
                                                package_ole.close()
                                                ole.close()
                                                return '.xls'
                                        except Exception as e:
                                            logger.warning(f"从package流的OLE2格式提取Workbook失败: {e}")
                                            # 如果提取失败，直接保存package数据
                                            extracted_path = os.path.join(save_dir, f"{ole_id}.xls")
                                            with open(extracted_path, 'wb') as f:
                                                f.write(package_data)
                                            logger.info(f"从package流提取数据，保存为.xls: {extracted_path}")
                                            ole.close()
                                            return '.xls'
                                    else:
                                        # 其他格式，根据ProgId判断
                                        if prog_id and 'excel' in prog_id.lower():
                                            extracted_path = os.path.join(save_dir, f"{ole_id}.xls")
                                            with open(extracted_path, 'wb') as f:
                                                f.write(package_data)
                                            logger.info(f"从package流提取数据，根据ProgId判断为.xls: {extracted_path}")
                                            ole.close()
                                            return '.xls'
                                        else:
                                            # 如果无法识别格式，但ProgId显示是Excel，尝试直接保存为.xls
                                            if prog_id and 'excel' in prog_id.lower():
                                                extracted_path = os.path.join(save_dir, f"{ole_id}.xls")
                                                with open(extracted_path, 'wb') as f:
                                                    f.write(package_data)
                                                logger.info(f"从package流提取数据，根据ProgId判断为.xls: {extracted_path}")
                                                ole.close()
                                                return '.xls'
                                            else:
                                                logger.warning(f"无法识别package流中的数据格式，ProgId: {prog_id}")
                                except Exception as e:
                                    logger.warning(f"从package流提取内容失败: {e}", exc_info=True)
                        
                        # 如果无法提取，但能判断类型，尝试其他方法
                        # 注意：只有在没有找到package流或package流提取失败时才执行
                        if detected_type and not stream_name_used:
                            logger.warning(f"✗ 检测到OLE2格式的{detected_type}文件，但无法提取标准流，尝试查找所有可能的流...")
                            
                            # 尝试查找所有可能的Excel流名（不区分大小写）
                            possible_streams = []
                            for stream_name in stream_list:
                                if isinstance(stream_name, tuple) and len(stream_name) > 0:
                                    stream_name_str = stream_name[0].lower()
                                else:
                                    stream_name_str = str(stream_name).lower()
                                
                                # 检查是否包含excel相关的关键词
                                if any(keyword in stream_name_str for keyword in ['workbook', 'book', 'excel', 'sheet', 'xls']):
                                    possible_streams.append(stream_name)
                            
                            if possible_streams:
                                # 尝试使用第一个可能的流
                                try:
                                    stream_to_try = possible_streams[0]
                                    extracted_data = ole.openstream(stream_to_try).read()
                                    extracted_path = os.path.join(save_dir, f"{ole_id}{detected_type}")
                                    with open(extracted_path, 'wb') as f:
                                        f.write(extracted_data)
                                    logger.info(f"✓ 使用备用流提取{detected_type}内容成功: {extracted_path}, 流名: {stream_to_try}")
                                    ole.close()
                                    return detected_type
                                except Exception as e:
                                    logger.warning(f"✗ 使用备用流提取失败: {e}")
                            
                            # 如果所有方法都失败，返回.bin
                            logger.warning(f"✗ 所有提取方法都失败，将保存为.bin格式: {ole_id}, 可用流: {all_streams_debug}")
                            ole.close()
                            return '.bin'  # 返回.bin，表示无法提取为标准格式
                        
                        # 如果既没有找到标准流，也没有找到package流，但ProgId显示是Excel，尝试查找所有流
                        if not stream_name_used and prog_id and 'excel' in prog_id.lower():
                            logger.warning(f"✗ 根据ProgId判断为Excel，但无法提取流，尝试查找所有可能的流...")
                            
                            # 尝试查找所有可能的Excel流
                            possible_streams = []
                            for stream_name in stream_list:
                                if isinstance(stream_name, tuple) and len(stream_name) > 0:
                                    stream_name_str = stream_name[0].lower()
                                else:
                                    stream_name_str = str(stream_name).lower()
                                
                                if any(keyword in stream_name_str for keyword in ['workbook', 'book', 'excel', 'sheet', 'xls']):
                                    possible_streams.append(stream_name)
                            
                            if possible_streams:
                                try:
                                    stream_to_try = possible_streams[0]
                                    extracted_data = ole.openstream(stream_to_try).read()
                                    extracted_path = os.path.join(save_dir, f"{ole_id}.xls")
                                    with open(extracted_path, 'wb') as f:
                                        f.write(extracted_data)
                                    logger.info(f"✓ 使用备用流提取.xls内容成功: {extracted_path}, 流名: {stream_to_try}")
                                    ole.close()
                                    return '.xls'
                                except Exception as e:
                                    logger.warning(f"✗ 使用备用流提取失败: {e}")
                            
                            logger.warning(f"✗ 所有提取方法都失败，将保存为.bin格式: {ole_id}, 可用流: {all_streams_debug}")
                            ole.close()
                            return '.bin'  # 返回.bin，表示无法提取为标准格式
                        
                    except Exception as e:
                        logger.warning(f"从OLE2格式提取内容失败: {e}", exc_info=True)
                    finally:
                        ole.close()
            except ImportError:
                logger.warning("olefile库未安装，无法提取OLE2格式中的内容。请安装: pip install olefile")
            except Exception as e:
                logger.debug(f"检测OLE2格式失败: {e}")
            
            # 如果无法提取，但能根据ProgId判断类型，返回对应扩展名
            if prog_id:
                prog_id_lower = prog_id.lower()
                if 'excel' in prog_id_lower:
                    logger.info(f"根据ProgId判断为Excel格式: {prog_id}")
                    return '.xls'
                elif 'word' in prog_id_lower:
                    logger.info(f"根据ProgId判断为Word格式: {prog_id}")
                    return '.doc'
                elif 'powerpoint' in prog_id_lower or 'ppt' in prog_id_lower:
                    logger.info(f"根据ProgId判断为PowerPoint格式: {prog_id}")
                    return '.ppt'
            
            # 如果无法判断，返回.bin
            logger.warning(f"无法识别OLE2格式的文件类型，返回.bin格式: {ole_id}")
            return '.bin'
        
        # 检查是否是旧版Excel格式（.xls）
        # 旧版Excel也是OLE2格式，但可能没有标准的OLE2签名
        if b'Microsoft Excel' in file_content[:1024] or b'Workbook' in file_content[:1024]:
            return '.xls'
        
        # 检查是否是Word格式
        if b'Microsoft Word' in file_content[:1024] or b'WordDocument' in file_content[:1024]:
            return '.doc'
        
        # 检查是否是PowerPoint格式
        if b'Microsoft PowerPoint' in file_content[:1024] or b'PowerPoint' in file_content[:1024]:
            return '.ppt'
        
        # 默认返回原始扩展名或.bin
        original_ext = os.path.splitext(original_path)[1].lower()
        if original_ext in ['.xlsx', '.xls', '.docx', '.doc', '.pptx', '.ppt']:
            return original_ext
        
        # 如果原始路径没有扩展名，但ProgId有信息，尝试根据ProgId判断
        if not original_ext or original_ext == '':
            if prog_id:
                prog_id_lower = prog_id.lower()
                if 'excel' in prog_id_lower:
                    return '.xls'
                elif 'word' in prog_id_lower:
                    return '.doc'
                elif 'powerpoint' in prog_id_lower or 'ppt' in prog_id_lower:
                    return '.ppt'
        
        return '.bin'
    
    @staticmethod
    def _generate_image_description(
        para_text: str,
        prev_paras_text: List[str],
        next_paras_text: List[str],
        section_title: str
    ) -> str:
        """
        生成图片描述（增强版）
        
        策略：
        1. 如果段落包含图片相关关键词，使用段落文本
        2. 如果前一段落包含图片相关关键词，使用前一段落文本
        3. 如果章节标题包含图片相关关键词，使用章节标题
        4. 否则，使用段落文本的前50字符
        """
        # 图片相关关键词（扩展版）
        image_keywords = [
            '图', '流程图', '示意图', '图片', '图表', '架构图', '时序图', 
            '用例图', '类图', '状态图', '活动图', '部署图', '组件图',
            'figure', 'image', 'diagram', 'chart', 'flowchart'
        ]
        
        # 策略1：检查当前段落
        if para_text:
            para_lower = para_text.lower()
            if any(keyword in para_lower for keyword in image_keywords):
                # 提取包含关键词的句子
                sentences = para_text.split('。') + para_text.split('.')
                for sentence in sentences:
                    sentence_lower = sentence.lower()
                    if any(keyword in sentence_lower for keyword in image_keywords):
                        return sentence.strip()[:100]
                return para_text[:100]
        
        # 策略2：检查前一段落
        if prev_paras_text:
            for prev_text in reversed(prev_paras_text):
                prev_lower = prev_text.lower()
                if any(keyword in prev_lower for keyword in image_keywords):
                    return f"位于段落：{prev_text[:80]}"
        
        # 策略3：检查章节标题
        if section_title:
            section_lower = section_title.lower()
            if any(keyword in section_lower for keyword in image_keywords):
                return f"{section_title}中的图片"
        
        # 策略4：使用当前段落文本（如果存在）
        if para_text:
            return f"位于段落：{para_text[:50]}"
        
        # 策略5：使用前一段落文本（如果存在）
        if prev_paras_text:
            return f"位于段落：{prev_paras_text[-1][:50]}"
        
        # 默认描述
        return "文档中的图片"
    
    @staticmethod
    def _format_table_as_text(table_data: Dict) -> str:
        """将表格数据格式化为文本描述（保留用于向后兼容）"""
        if not table_data.get("headers"):
            return ""
        
        # 构建表格文本
        text = "表格：\n"
        text += " | ".join(table_data["headers"]) + "\n"
        text += "-" * (len(" | ".join(table_data["headers"]))) + "\n"
        
        for row in table_data.get("rows", []):
            text += " | ".join(row) + "\n"
        
        return text
    
    @staticmethod
    def _format_table_as_markdown(table_data: Dict) -> str:
        """
        将表格数据格式化为标准Markdown表格
        
        Args:
            table_data: {
                "headers": ["列1", "列2", "列3"],
                "rows": [["值1", "值2", "值3"], ...],
                "row_count": int,
                "col_count": int
            }
        
        Returns:
            标准Markdown格式的表格字符串
        """
        if not table_data.get("headers"):
            return ""
        
        headers = table_data["headers"]
        rows = table_data.get("rows", [])
        
        # 构建标准Markdown表格
        # 表头行
        markdown = "| " + " | ".join(str(header) for header in headers) + " |\n"
        # 分隔行（标准Markdown表格格式）
        markdown += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        
        # 数据行
        for row in rows:
            # 确保行数据长度与表头一致
            row_data = row[:len(headers)] if row else []
            if len(row_data) < len(headers):
                row_data.extend([""] * (len(headers) - len(row_data)))
            # 转义表格中的管道符，避免破坏表格结构
            escaped_row = [str(cell).replace("|", "\\|") for cell in row_data]
            markdown += "| " + " | ".join(escaped_row) + " |\n"
        
        return markdown
    
    @staticmethod
    def _extract_images_from_document(doc: Document, document_id: str = None, file_path: str = None) -> List[Dict]:
        """
        提取文档中的图片并保存到文件系统
        
        增强功能：
        1. 更准确的关系ID映射（支持多种路径格式）
        2. 图片元数据提取（大小、格式等）
        3. 匹配统计信息
        4. 增强的上下文信息（多段落、章节标题、相对位置）
        """
        images = []
        match_stats = {
            "total_images": 0,
            "with_rel_id": 0,
            "without_rel_id": 0,
            "matched_by_rel_id": 0,
            "matched_by_drawing": 0,
            "matched_by_keyword": 0,
            "unmatched": 0
        }
        try:
            from docx.oxml.ns import qn
            import zipfile
            import shutil
            
            image_counter = 0
            
            # 创建图片保存目录
            if document_id and file_path:
                # 使用document_id创建子目录
                image_dir = os.path.abspath(f"uploads/extracted_images/{document_id}")
                os.makedirs(image_dir, exist_ok=True)
            else:
                image_dir = os.path.abspath("uploads/extracted_images/temp")
                os.makedirs(image_dir, exist_ok=True)
            
            # 从docx文件中提取图片（docx是zip格式）
            # 建立关系ID到图片文件的映射
            rel_id_to_image_file = {}  # {rId: image_file_path}
            image_file_to_rel_id = {}  # {image_file_path: rId}
            
            if file_path and os.path.exists(file_path):
                try:
                    # 首先读取关系文件，建立关系ID到图片文件的映射
                    with zipfile.ZipFile(file_path, 'r') as zip_file:
                        # 读取document.xml.rels文件，建立关系ID到图片文件的映射
                        try:
                            rels_file = zip_file.read('word/_rels/document.xml.rels')
                            import xml.etree.ElementTree as ET
                            rels_root = ET.fromstring(rels_file)
                            
                            # 解析关系文件，找到所有图片关系
                            for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                                rel_type = rel.get('Type', '')
                                target = rel.get('Target', '')
                                rel_id = rel.get('Id', '')
                                
                                # 如果是图片关系
                                if 'image' in rel_type.lower() or target.startswith('media/'):
                                    # 标准化路径
                                    if target.startswith('media/'):
                                        image_file_path = f"word/{target}"
                                    else:
                                        image_file_path = f"word/media/{target}"
                                    
                                    rel_id_to_image_file[rel_id] = image_file_path
                                    image_file_to_rel_id[image_file_path] = rel_id
                                    logger.debug(f"建立关系映射: rId={rel_id} -> {image_file_path}")
                        except Exception as e:
                            logger.warning(f"读取关系文件失败: {e}，将使用备用方法")
                        
                        # 直接解析文档结构，找到所有图片出现的位置
                        # 不需要预先查找所有图片文件，而是遍历文档时遇到图片就提取
                        
                        logger.info(f"📋 建立关系映射: {len(rel_id_to_image_file)} 个图片关系")
                        
                        # 现在遍历文档段落，按照图片在文档中出现的顺序分配image_id
                        # 构建章节标题映射（用于上下文增强）
                        section_titles = []
                        current_section_title = ""
                        for para_idx, para in enumerate(doc.paragraphs):
                            if para.style.name.startswith('Heading'):
                                current_section_title = para.text.strip()
                                section_titles.append((para_idx, current_section_title))
                        
                        # 计算文档总段落数（用于相对位置计算）
                        total_paragraphs = len(doc.paragraphs)
                        
                        # 遍历所有段落，找到图片出现的位置（按文档顺序）
                        for para_idx, paragraph in enumerate(doc.paragraphs):
                            para_text = paragraph.text.strip()
                            
                            # 获取前后多个段落的文本作为上下文
                            prev_paras_text = []
                            next_paras_text = []
                            
                            for i in range(max(0, para_idx - 2), para_idx):
                                if i < len(doc.paragraphs):
                                    prev_text = doc.paragraphs[i].text.strip()
                                    if prev_text:
                                        prev_paras_text.append(prev_text)
                            
                            for i in range(para_idx + 1, min(para_idx + 3, len(doc.paragraphs))):
                                if i < len(doc.paragraphs):
                                    next_text = doc.paragraphs[i].text.strip()
                                    if next_text:
                                        next_paras_text.append(next_text)
                            
                            # 获取最近的章节标题
                            nearest_section_title = ""
                            for section_idx, section_title in reversed(section_titles):
                                if section_idx <= para_idx:
                                    nearest_section_title = section_title
                                    break
                            
                            # 计算相对位置
                            relative_position = para_idx / total_paragraphs if total_paragraphs > 0 else 0.0
                            
                            for run in paragraph.runs:
                                # 检查run中是否有图片
                                blips = run._element.xpath('.//a:blip')
                                
                                if blips:
                                    # 通过关系ID精确匹配图片
                                    for blip in blips:
                                        embed_id = blip.get(qn('r:embed'))
                                        link_id = blip.get(qn('r:link'))
                                        rel_id = embed_id or link_id
                                        
                                        if rel_id and rel_id in rel_id_to_image_file:
                                            # 找到图片，按文档顺序分配image_id
                                            img_file = rel_id_to_image_file[rel_id]
                                            
                                            # 检查图片文件是否存在
                                            if img_file in zip_file.namelist():
                                                image_counter += 1
                                                image_id = f"image_{image_counter}"
                                                
                                                # 获取文件扩展名
                                                ext = os.path.splitext(img_file)[1] or '.png'
                                                file_name = os.path.basename(img_file)
                                                
                                                # 生成描述
                                                description = WordDocumentService._generate_image_description(
                                                    para_text, prev_paras_text, next_paras_text, nearest_section_title
                                                )
                                                
                                                # 保存图片文件
                                                saved_image_path = os.path.join(image_dir, f"{image_id}{ext}")
                                                with zip_file.open(img_file) as source, open(saved_image_path, 'wb') as target:
                                                    shutil.copyfileobj(source, target)
                                                
                                                # 获取文件大小和格式
                                                file_size = os.path.getsize(saved_image_path) if os.path.exists(saved_image_path) else 0
                                                file_format = ext[1:].upper() if ext else 'UNKNOWN'  # 去掉点号，转为大写
                                                
                                                # 获取相对路径
                                                relative_path = f"extracted_images/{document_id}/{image_id}{ext}" if document_id else f"extracted_images/temp/{image_id}{ext}"
                                                
                                                # 构建完整的图片数据
                                                images.append({
                                                    "image_id": image_id,
                                                    "position": para_idx,
                                                    "description": description,
                                                    "file_path": saved_image_path,
                                                    "relative_path": relative_path,
                                                    "file_name": file_name,
                                                    "rel_id": rel_id,
                                                    "file_size": file_size,  # 添加文件大小
                                                    "file_format": file_format,  # 添加文件格式
                                                    "context": para_text[:300] if para_text else "",
                                                    "prev_context": " | ".join(prev_paras_text[:2])[:200] if prev_paras_text else "",
                                                    "next_context": " | ".join(next_paras_text[:2])[:200] if next_paras_text else "",
                                                    "section_title": nearest_section_title,
                                                    "relative_position": relative_position,
                                                    "match_method": "rel_id",
                                                    "match_confidence": 1.0
                                                })
                                                
                                                match_stats["matched_by_rel_id"] += 1
                                                match_stats["with_rel_id"] += 1
                                                logger.info(f"✅ 图片 {image_id} 通过关系ID匹配到段落 {para_idx} (rel_id: {rel_id}, 文件: {file_name}, 章节: {nearest_section_title[:30]})")
                                                
                                                # 注意：不删除，允许同一图片文件多次出现
                                                # 继续处理下一个blip（一个段落可能有多张图片）
                                            else:
                                                logger.warning(f"⚠️ 图片文件不存在: {img_file} (rel_id: {rel_id})")
                                
                                
                except Exception as e:
                    logger.warning(f"从zip文件提取图片失败: {e}", exc_info=True)
            
            # 注意：图片编号已经在zip_file块内按文档顺序分配完成
            # 所有图片（包括未匹配的）都已经在zip_file块内处理并分配了image_id
            
            # 更新统计信息
            match_stats["total_images"] = len(images)
            match_stats["unmatched"] = sum(1 for img in images if img.get("position") == -1)
            
            # 输出匹配统计信息
            logger.info(f"📊 图片匹配统计: 总数={match_stats['total_images']}, "
                       f"有rel_id={match_stats['with_rel_id']}, "
                       f"无rel_id={match_stats['without_rel_id']}, "
                       f"关系ID匹配={match_stats['matched_by_rel_id']}, "
                       f"drawing匹配={match_stats['matched_by_drawing']}, "
                       f"未匹配={match_stats['unmatched']}")
            
            # 将匹配统计信息添加到第一张图片的元数据中（用于后续分析）
            if images:
                images[0]["_match_stats"] = match_stats
            
            logger.info(f"从文档中提取到 {len(images)} 张图片，已保存到 {image_dir}")
        except Exception as e:
            logger.warning(f"提取图片时出错: {e}", exc_info=True)
        
        return images
    
    @staticmethod
    def _split_by_sections(structured_content: List[Dict], max_tokens: int = 8000) -> List[Dict]:
        """按章节分块"""
        sections = []
        current_section = None
        
        for item in structured_content:
            if item["type"] == "heading":
                level = item.get("level", 1)
                if level == 1:
                    # 一级标题：创建新章节
                    if current_section and current_section.get("token_count", 0) > 0:
                        sections.append(current_section)
                    
                    # 创建新章节
                    current_section = {
                        "section_id": f"section_{len(sections)}",
                        "title": item["text"],
                        "level": level,
                        "content": "",  # 一级标题不重复添加到content中
                        "token_count": 0,
                        "images": [],
                        "links": item.get("links", []),
                        "tables": []
                    }
                else:
                    # 子标题（level > 1）：保留在父章节中，使用Markdown格式
                    if current_section is None:
                        # 如果没有章节，创建默认章节
                        current_section = {
                            "section_id": "section_0",
                            "title": "概述",
                            "level": 1,
                            "content": "",
                            "token_count": 0,
                            "images": [],
                            "links": [],
                            "tables": []
                        }
                    
                    # 使用Markdown标题格式（##, ###, ####等）
                    heading_markdown = "#" * (level + 1)  # level=2 -> ##, level=3 -> ###
                    heading_text = item["text"]
                    heading_content = f"{heading_markdown} {heading_text}\n\n"
                    
                    # 检查是否需要分割（超过最大 token 数）
                    heading_tokens = WordDocumentService._estimate_tokens(heading_content)
                    if current_section["token_count"] + heading_tokens > max_tokens:
                        # 保存当前章节
                        sections.append(current_section)
                        
                        # 创建新章节（子章节）
                        current_section = {
                            "section_id": f"section_{len(sections)}",
                            "title": current_section["title"] + "（续）",
                            "level": current_section["level"],
                            "content": current_section["title"] + "\n\n",
                            "token_count": WordDocumentService._estimate_tokens(current_section["title"]),
                            "images": [],
                            "links": [],
                            "tables": []
                        }
                    
                    # 添加子标题到内容中
                    current_section["content"] += heading_content
                    current_section["token_count"] += heading_tokens
                    
                    # 处理子标题的链接
                    if item.get("links"):
                        current_section["links"].extend(item.get("links", []))
            else:
                # 添加到当前章节
                if current_section is None:
                    # 如果没有章节，创建默认章节
                    current_section = {
                        "section_id": "section_0",
                        "title": "概述",
                        "level": 1,
                        "content": "",
                        "token_count": 0,
                        "images": [],
                        "links": [],
                        "tables": []
                    }
                
                # 处理表格类型
                if item["type"] == "table":
                    # 表格直接添加到当前章节
                    current_section["tables"].append(item["data"])
                    # 表格文本也添加到内容中
                    table_text = item.get("text", "")
                    if table_text:
                        item_tokens = WordDocumentService._estimate_tokens(table_text)
                        # 检查是否需要分割
                        if current_section["token_count"] + item_tokens > max_tokens:
                            # 保存当前章节
                            sections.append(current_section)
                            
                            # 创建新章节（子章节）
                            current_section = {
                                "section_id": f"section_{len(sections)}",
                                "title": current_section["title"] + "（续）",
                                "level": current_section["level"],
                                "content": current_section["title"] + "\n",
                                "token_count": WordDocumentService._estimate_tokens(current_section["title"]),
                                "images": [],
                                "links": [],
                                "tables": []
                            }
                        current_section["content"] += table_text + "\n"
                        current_section["token_count"] += item_tokens
                else:
                    # 处理段落、图片等其他类型
                    item_text = item.get("text", "")
                    item_tokens = WordDocumentService._estimate_tokens(item_text) if item_text else 0
                    
                    # 检查是否需要分割（超过最大 token 数）
                    if item_tokens > 0 and current_section["token_count"] + item_tokens > max_tokens:
                        # 保存当前章节
                        sections.append(current_section)
                        
                        # 创建新章节（子章节）
                        current_section = {
                            "section_id": f"section_{len(sections)}",
                            "title": current_section["title"] + "（续）",
                            "level": current_section["level"],
                            "content": current_section["title"] + "\n",  # 保留标题
                            "token_count": WordDocumentService._estimate_tokens(current_section["title"]),
                            "images": [],
                            "links": [],
                            "tables": []
                        }
                    
                    # 添加内容（确保段落之间有适当的空行）
                    if item_text:
                        current_section["content"] += item_text + "\n\n"
                        current_section["token_count"] += item_tokens
                    
                    # 处理图片
                    if item.get("images"):
                        if "images" not in current_section:
                            current_section["images"] = []
                        current_section["images"].extend(item["images"])
        
        # 添加最后一个章节
        if current_section and current_section.get("token_count", 0) > 0:
            sections.append(current_section)
        
        return sections
    
    @staticmethod
    def _split_by_sections_with_strategy(
        structured_content: List[Dict], 
        strategy: str = "level_1",
        max_tokens: int = 8000
    ) -> List[Dict]:
        """
        根据策略进行分块
        
        策略：
        - level_1: 按一级标题分块（默认）
        - level_2: 按二级标题分块
        - level_3: 按三级标题分块
        - level_4: 按四级标题分块
        - level_5: 按五级标题分块
        - fixed_token: 按固定token数分块
        - no_split: 不分块（整个文档作为一个块）
        
        返回：
        每个块包含 start_index, end_index, content 字段
        """
        if strategy == "no_split":
            # 不分块：整个文档作为一个块
            return WordDocumentService._split_no_split(structured_content, max_tokens)
        elif strategy == "fixed_token":
            # 按固定 token 数分块
            return WordDocumentService._split_by_fixed_tokens(structured_content, max_tokens)
        elif strategy == "level_2":
            # 按二级标题分块
            return WordDocumentService._split_by_heading_level(structured_content, 2, max_tokens)
        elif strategy == "level_3":
            # 按三级标题分块
            return WordDocumentService._split_by_heading_level(structured_content, 3, max_tokens)
        elif strategy == "level_4":
            # 按四级标题分块
            return WordDocumentService._split_by_heading_level(structured_content, 4, max_tokens)
        elif strategy == "level_5":
            # 按五级标题分块
            return WordDocumentService._split_by_heading_level(structured_content, 5, max_tokens)
        else:  # level_1 或默认
            # 按一级标题分块
            return WordDocumentService._split_by_heading_level(structured_content, 1, max_tokens)
    
    @staticmethod
    def _split_no_split(structured_content: List[Dict], max_tokens: int = 8000) -> List[Dict]:
        """不分块：整个文档作为一个块"""
        content = ""
        total_tokens = 0
        images = []
        tables = []
        links = []
        
        for idx, item in enumerate(structured_content):
            item_text = item.get("text", "")
            if item["type"] == "heading":
                level = item.get("level", 1)
                heading_markdown = "#" * level
                content += f"{heading_markdown} {item_text}\n\n"
            elif item["type"] == "table":
                content += item_text + "\n\n"
                tables.append(item.get("data", []))
            else:
                if item_text:
                    content += item_text + "\n\n"
            
            if item.get("images"):
                images.extend(item["images"])
            if item.get("links"):
                links.extend(item["links"])
            
            total_tokens += WordDocumentService._estimate_tokens(item_text)
        
        # 确定标题
        title = "全文档"
        for item in structured_content:
            if item["type"] == "heading" and item.get("level", 1) == 1:
                title = item.get("text", "全文档")
                break
        
        return [{
            "section_id": "chunk_1",
            "title": title,
            "level": 1,
            "content": content.strip(),
            "token_count": total_tokens,
            "start_index": 0,
            "end_index": len(structured_content),
            "images": images,
            "tables": tables,
            "links": links
        }]
    
    @staticmethod
    def _split_by_fixed_tokens(structured_content: List[Dict], max_tokens: int = 8000) -> List[Dict]:
        """按固定 token 数分块"""
        sections = []
        current_section = {
            "section_id": "chunk_1",
            "title": "段落 1",
            "level": 1,
            "content": "",
            "token_count": 0,
            "start_index": 0,
            "end_index": 0,
            "images": [],
            "tables": [],
            "links": []
        }
        
        for idx, item in enumerate(structured_content):
            item_text = item.get("text", "")
            item_tokens = WordDocumentService._estimate_tokens(item_text) if item_text else 0
            
            # 构建内容
            item_content = ""
            if item["type"] == "heading":
                level = item.get("level", 1)
                heading_markdown = "#" * level
                item_content = f"{heading_markdown} {item_text}\n\n"
            elif item["type"] == "table":
                item_content = item_text + "\n\n"
            else:
                if item_text:
                    item_content = item_text + "\n\n"
            
            # 检查是否需要分块
            if current_section["token_count"] + item_tokens > max_tokens and current_section["token_count"] > 0:
                # 保存当前块
                current_section["end_index"] = idx
                sections.append(current_section)
                
                # 创建新块
                chunk_num = len(sections) + 1
                current_section = {
                    "section_id": f"chunk_{chunk_num}",
                    "title": f"段落 {chunk_num}",
                    "level": 1,
                    "content": "",
                    "token_count": 0,
                    "start_index": idx,
                    "end_index": 0,
                    "images": [],
                    "tables": [],
                    "links": []
                }
            
            # 添加内容
            current_section["content"] += item_content
            current_section["token_count"] += item_tokens
            
            if item.get("images"):
                current_section["images"].extend(item["images"])
            if item.get("links"):
                current_section["links"].extend(item["links"])
            if item["type"] == "table":
                current_section["tables"].append(item.get("data", []))
        
        # 添加最后一个块
        if current_section["token_count"] > 0:
            current_section["end_index"] = len(structured_content)
            sections.append(current_section)
        
        return sections
    
    @staticmethod
    def _split_by_heading_level(
        structured_content: List[Dict], 
        split_level: int = 1,
        max_tokens: int = 8000
    ) -> List[Dict]:
        """按指定标题级别分块"""
        sections = []
        current_section = None
        current_start_index = 0
        
        for idx, item in enumerate(structured_content):
            if item["type"] == "heading" and item.get("level", 1) <= split_level:
                # 遇到分割标题，保存之前的块
                if current_section and current_section.get("token_count", 0) > 0:
                    current_section["end_index"] = idx
                    sections.append(current_section)
                
                # 创建新块
                current_start_index = idx
                current_section = {
                    "section_id": f"chunk_{len(sections) + 1}",
                    "title": item["text"],
                    "level": item.get("level", 1),
                    "content": "",
                    "token_count": 0,
                    "start_index": idx,
                    "end_index": 0,
                    "images": [],
                    "tables": [],
                    "links": item.get("links", [])
                }
            else:
                # 添加到当前块
                if current_section is None:
                    # 如果没有块，创建默认块
                    current_section = {
                        "section_id": "chunk_1",
                        "title": "概述",
                        "level": 1,
                        "content": "",
                        "token_count": 0,
                        "start_index": 0,
                        "end_index": 0,
                        "images": [],
                        "tables": [],
                        "links": []
                    }
                
                item_text = item.get("text", "")
                item_tokens = WordDocumentService._estimate_tokens(item_text) if item_text else 0
                
                # 构建内容
                if item["type"] == "heading":
                    level = item.get("level", 1)
                    heading_markdown = "#" * level
                    current_section["content"] += f"{heading_markdown} {item_text}\n\n"
                elif item["type"] == "table":
                    current_section["content"] += item_text + "\n\n"
                    current_section["tables"].append(item.get("data", []))
                else:
                    if item_text:
                        current_section["content"] += item_text + "\n\n"
                
                current_section["token_count"] += item_tokens
                
                if item.get("images"):
                    current_section["images"].extend(item["images"])
                if item.get("links"):
                    current_section["links"].extend(item["links"])
        
        # 添加最后一个块
        if current_section and current_section.get("token_count", 0) > 0:
            current_section["end_index"] = len(structured_content)
            sections.append(current_section)
        
        return sections
    
    @staticmethod
    def _extract_formatted_text(paragraph) -> str:
        """
        提取段落文本，保留格式信息（转换为Markdown）
        
        Args:
            paragraph: docx段落对象
        
        Returns:
            格式化的文本（Markdown格式）
        """
        formatted_text = ""
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # 处理加粗
            if run.bold:
                text = f"**{text}**"
            
            # 处理斜体
            if run.italic:
                text = f"*{text}*"
            
            # 处理下划线（Markdown不直接支持，使用HTML）
            if run.underline:
                text = f"<u>{text}</u>"
            
            # 处理删除线
            if hasattr(run, 'strike') and run.strike:
                text = f"~~{text}~~"
            
            formatted_text += text
        
        return formatted_text.strip()
    
    @staticmethod
    def _estimate_tokens(text: str) -> int:
        """估算文本的 token 数（中文通常 1 token ≈ 2 字符）"""
        return len(text) // 2
    
    @staticmethod
    def _extract_base_name(document_name: str) -> str:
        """
        从文档名称中提取基础标识（去除版本号）
        
        Args:
            document_name: 文档名称，例如 "产业项目-项目里程碑管理-软件需求规格说明书-20230731- V1"
        
        Returns:
            基础标识，例如 "产业项目-项目里程碑管理-软件需求规格说明书-20230731"
        """
        import re
        
        # 支持多种版本号格式
        patterns = [
            r'\s*-\s*V\d+$',           # " - V1"
            r'\s*-\s*v\d+$',           # " - v1"
            r'\s*版本\d+$',             # " 版本1"
            r'\s*Version\s*\d+$',      # " Version 1"
            r'\s*version\s*\d+$',      # " version 1"
        ]
        
        base_name = document_name
        for pattern in patterns:
            base_name = re.sub(pattern, '', base_name, flags=re.IGNORECASE)
        
        return base_name.strip()
    
    @staticmethod
    def _extract_version(document_name: str) -> tuple[str, int]:
        """
        从文档名称中提取版本号
        
        Args:
            document_name: 文档名称
        
        Returns:
            (version_string, version_number) 例如 ("V1", 1)
        """
        import re
        
        version_match = re.search(r'V(\d+)', document_name, re.IGNORECASE)
        if version_match:
            version_num = int(version_match.group(1))
            return f"V{version_num}", version_num
        
        # 如果找不到，返回默认值
        return "V1", 1
    
    @staticmethod
    def _sanitize_group_id(name: str) -> str:
        """
        清理 group_id，只保留字母数字、破折号和下划线
        
        Graphiti 要求 group_id 只能包含 alphanumeric characters, dashes, or underscores
        
        Args:
            name: 原始名称（可能包含中文、特殊字符等）
        
        Returns:
            清理后的名称（只包含字母数字、破折号、下划线）
        """
        import re
        
        # 将中文字符和其他特殊字符替换为下划线
        # 只保留字母数字、破折号、下划线
        sanitized = re.sub(r'[^a-zA-Z0-9\-_]', '_', name)
        
        # 将连续的下划线替换为单个下划线
        sanitized = re.sub(r'_+', '_', sanitized)
        
        # 去除开头和结尾的下划线
        sanitized = sanitized.strip('_')
        
        # 如果清理后为空，使用默认值
        if not sanitized:
            sanitized = "document"
        
        # 限制长度（避免过长）
        if len(sanitized) > 50:
            sanitized = sanitized[:50]
        
        return sanitized
    
    @staticmethod
    def _build_section_content(section: Dict, doc_data: Dict, section_idx: int, document_id: str = None, upload_id: int = None) -> str:
        """
        构建章节内容（1:1对应原始文档，不添加额外描述）
        
        只包含原始文档中的内容：
        - 标题和段落文字
        - 图片：![alt](url)（alt使用原始图注或空）
        - 表格：只保留表格内容，不添加标题（除非原始文档有）
        - 嵌入文档：简单占位符
        - 链接：保留在原始位置
        
        注意：此方法直接遍历整个structured_content，按原始顺序输出，确保1:1对应。
        如果文档有多个章节，每个章节会输出一次，但内容不会重复。
        """
        # 将document_id添加到doc_data中，以便在构建内容时使用
        if document_id and 'document_id' not in doc_data:
            doc_data['document_id'] = document_id
        
        # 找到当前章节在structured_content中的范围
        section_start_idx = None
        section_end_idx = None
        section_title_found = False  # 标记是否在原始文档中找到该章节标题
        
        # 找到当前章节的起始位置（一级标题）
        for idx, item in enumerate(doc_data.get("structured_content", [])):
            if item.get("type") == "heading" and item.get("level", 1) == 1:
                # 如果找到了匹配的标题，记录起始位置
                if item.get("text") == section.get("title"):
                    section_start_idx = idx
                    section_title_found = True
                    # 找到下一个一级标题作为结束位置
                    for next_idx in range(idx + 1, len(doc_data.get("structured_content", []))):
                        next_item = doc_data["structured_content"][next_idx]
                        if next_item.get("type") == "heading" and next_item.get("level", 1) == 1:
                            section_end_idx = next_idx
                            break
                    break
        
        # 如果没有找到匹配的标题（可能是系统生成的"概述"），不输出任何内容
        # 这样可以避免重复输出整个文档
        if not section_title_found:
            # 如果标题是系统生成的"概述"，且原始文档中没有一级标题，则输出整个文档
            # 但只输出一次（第一个章节）
            if section.get("title") == "概述" and section_idx == 0:
                # 检查原始文档是否真的没有一级标题
                has_level1_heading = any(
                    item.get("type") == "heading" and item.get("level", 1) == 1
                    for item in doc_data.get("structured_content", [])
                )
                if not has_level1_heading:
                    # 原始文档确实没有一级标题，输出整个文档
                    section_start_idx = 0
                    section_end_idx = len(doc_data.get("structured_content", []))
                else:
                    # 原始文档有一级标题，但当前章节标题不匹配，不输出
                    return ""
            else:
                # 其他情况，不输出
                return ""
        
        # 如果找到了起始位置，构建内容
        if section_start_idx is None:
            return ""
        
        # 如果没有找到结束位置，使用文档末尾
        if section_end_idx is None:
            section_end_idx = len(doc_data.get("structured_content", []))
        
        # 构建章节内容（按照structured_content的顺序）
        # 只有当原始文档中存在该章节标题时，才输出标题
        # 注意：标题文本应该包含原始文档中的完整文本（包括数字前缀，如"1 项目里程碑管理"）
        content = ""
        if section_title_found:
            # 从原始文档中获取完整的标题文本（包含数字前缀）
            original_title_text = ""
            for idx, item in enumerate(doc_data.get("structured_content", [])):
                if idx == section_start_idx and item.get("type") == "heading":
                    original_title_text = item.get("text", section.get("title", ""))
                    break
            # 如果找到了原始标题文本，使用它；否则使用section中的title
            title_to_use = original_title_text if original_title_text else section.get("title", "")
            content = f"# {title_to_use}\n\n"
        
        # 按照structured_content的顺序构建内容
        for idx in range(section_start_idx, min(section_end_idx, len(doc_data.get("structured_content", [])))):
            if idx >= len(doc_data.get("structured_content", [])):
                break
            item = doc_data["structured_content"][idx]
            
            # 跳过一级标题（已经在上面添加了）
            if item.get("type") == "heading" and item.get("level", 1) == 1:
                continue
            
            # 处理子标题（level > 1）
            if item.get("type") == "heading" and item.get("level", 1) > 1:
                level = item.get("level", 2)
                heading_markdown = "#" * (level + 1)  # level=2 -> ##, level=3 -> ###
                heading_text = item.get("text", "")
                content += f"{heading_markdown} {heading_text}\n\n"
            
            # 处理段落
            elif item.get("type") == "paragraph":
                paragraph_text = item.get("text", "")
                if paragraph_text:
                    content += f"{paragraph_text}\n\n"
                
                # 如果段落有关联的图片，立即插入（保留原始位置）
                # 只保留图片链接，不添加额外描述
                if item.get("images"):
                    for image in item["images"]:
                        image_id = image.get('image_id', '')
                        relative_path = image.get('relative_path', '')
                        
                        # 尝试获取原始图注（如果有），否则使用空alt text
                        # 注意：当前代码中没有提取原始图注的逻辑，暂时使用空alt text
                        alt_text = ""  # 可以后续扩展提取原始图注的逻辑
                        
                        if relative_path and doc_data.get('document_id'):
                            document_id = doc_data.get('document_id')
                            image_url = f"/api/word-document/{document_id}/images/{image_id}"
                            content += f"![{alt_text}]({image_url})\n\n"
            
            # 处理表格（保留原始位置）
            # 不添加系统生成的标题，只保留表格内容
            elif item.get("type") == "table":
                table_data = item.get("data", {})
                
                # 注意：当前代码中没有提取原始表格标题的逻辑
                # 如果后续需要支持原始表格标题，需要在这里添加提取逻辑
                # 暂时只保留表格内容，不添加任何标题
                
                # 使用标准Markdown表格格式
                content += WordDocumentService._format_table_as_markdown(table_data) + "\n\n"
            
            # 处理image_only类型（单独的图片）
            elif item.get("type") == "image_only":
                if item.get("images"):
                    for image in item["images"]:
                        image_id = image.get('image_id', '')
                        relative_path = image.get('relative_path', '')
                        
                        # 只保留图片链接，不添加额外描述
                        alt_text = ""  # 可以后续扩展提取原始图注的逻辑
                        
                        if relative_path and doc_data.get('document_id'):
                            document_id = doc_data.get('document_id')
                            image_url = f"/api/word-document/{document_id}/images/{image_id}"
                            content += f"![{alt_text}]({image_url})\n\n"
            
            # 处理OLE对象（嵌入文档）
            # 生成包含完整信息的Markdown格式，与"需求管理"的显示方式一致
            if item.get("ole_objects"):
                for ole in item.get("ole_objects", []):
                    ole_id = ole.get('ole_id', '')
                    ole_name = ole.get('name', '嵌入文档')
                    ole_type = ole.get('type', '嵌入对象')
                    document_id = doc_data.get('document_id', '')
                    
                    # 如果upload_id存在，使用document-upload API；否则使用word-document API
                    if upload_id and ole_id:
                        preview_url = f"/api/document-upload/{upload_id}/ole/{ole_id}?view=preview"
                        download_url = f"/api/document-upload/{upload_id}/ole/{ole_id}?view=download"
                        # 生成包含文件名、类型、查看/下载链接的Markdown
                        content += f"[嵌入文档: {ole_name} ({ole_type})]({preview_url})\n"
                        content += f"[查看]({preview_url}) | [下载]({download_url})\n\n"
                    elif document_id and ole_id:
                        # 兼容旧格式（使用word-document API）
                        preview_url = f"/api/word-document/{document_id}/ole/{ole_id}?view=preview"
                        download_url = f"/api/word-document/{document_id}/ole/{ole_id}?view=download"
                        content += f"[嵌入文档: {ole_name} ({ole_type})]({preview_url})\n"
                        content += f"[查看]({preview_url}) | [下载]({download_url})\n\n"
                    else:
                        content += "[嵌入文档]\n\n"
        
        # 注意：不再添加章节末尾的链接列表，链接保留在原始位置（段落中）
        
        return content
    
    @staticmethod
    def _build_content_from_items(items: List[Dict], doc_data: Dict, document_id: str = None, upload_id: int = None) -> str:
        """
        从structured_content的items列表构建内容（用于封面页、表格等）
        
        这个方法用于构建第一个一级标题之前的内容，确保1:1对应原始文档
        
        Args:
            items: structured_content的子列表
            doc_data: 完整的文档数据
            document_id: 文档ID（用于构建图片和OLE对象的URL）
        """
        if document_id and 'document_id' not in doc_data:
            doc_data['document_id'] = document_id
        
        content = ""
        
        for item in items:
            # 处理段落
            if item.get("type") == "paragraph":
                paragraph_text = item.get("text", "")
                if paragraph_text:
                    content += f"{paragraph_text}\n\n"
                
                # 如果段落有关联的图片，立即插入（保留原始位置）
                if item.get("images"):
                    for image in item["images"]:
                        image_id = image.get('image_id', '')
                        relative_path = image.get('relative_path', '')
                        alt_text = ""  # 可以后续扩展提取原始图注的逻辑
                        
                        if relative_path and doc_data.get('document_id'):
                            document_id = doc_data.get('document_id')
                            image_url = f"/api/word-document/{document_id}/images/{image_id}"
                            content += f"![{alt_text}]({image_url})\n\n"
            
            # 处理标题（可能是封面页的标题，不是一级标题）
            elif item.get("type") == "heading":
                level = item.get("level", 1)
                heading_text = item.get("text", "")
                if heading_text:
                    # 使用对应的Markdown标题级别
                    heading_markdown = "#" * (level + 1)  # level=1 -> ##, level=2 -> ###
                    content += f"{heading_markdown} {heading_text}\n\n"
            
            # 处理表格（保留原始位置）
            elif item.get("type") == "table":
                table_data = item.get("data", {})
                # 如果原始Word文档有表格标题，则保留
                if table_data.get("caption"):
                    content += f"### {table_data['caption']}\n\n"
                # 使用标准Markdown表格格式
                content += WordDocumentService._format_table_as_markdown(table_data) + "\n\n"
            
            # 处理image_only类型（单独的图片）
            elif item.get("type") == "image_only":
                if item.get("images"):
                    for image in item["images"]:
                        image_id = image.get('image_id', '')
                        relative_path = image.get('relative_path', '')
                        alt_text = ""
                        
                        if relative_path and doc_data.get('document_id'):
                            document_id = doc_data.get('document_id')
                            image_url = f"/api/word-document/{document_id}/images/{image_id}"
                            content += f"![{alt_text}]({image_url})\n\n"
            
            # 处理OLE对象（嵌入文档）
            # 生成包含完整信息的Markdown格式，与"需求管理"的显示方式一致
            if item.get("ole_objects"):
                for ole in item.get("ole_objects", []):
                    ole_id = ole.get('ole_id', '')
                    ole_name = ole.get('name', '嵌入文档')
                    ole_type = ole.get('type', '嵌入对象')
                    document_id = doc_data.get('document_id', '')
                    
                    # 如果upload_id存在，使用document-upload API；否则使用word-document API
                    if upload_id and ole_id:
                        preview_url = f"/api/document-upload/{upload_id}/ole/{ole_id}?view=preview"
                        download_url = f"/api/document-upload/{upload_id}/ole/{ole_id}?view=download"
                        # 生成包含文件名、类型、查看/下载链接的Markdown
                        content += f"[嵌入文档: {ole_name} ({ole_type})]({preview_url})\n"
                        content += f"[查看]({preview_url}) | [下载]({download_url})\n\n"
                    elif document_id and ole_id:
                        # 兼容旧格式（使用word-document API）
                        preview_url = f"/api/word-document/{document_id}/ole/{ole_id}?view=preview"
                        download_url = f"/api/word-document/{document_id}/ole/{ole_id}?view=download"
                        content += f"[嵌入文档: {ole_name} ({ole_type})]({preview_url})\n"
                        content += f"[查看]({preview_url}) | [下载]({download_url})\n\n"
                    else:
                        content += "[嵌入文档]\n\n"
        
        return content.strip()
    
    @staticmethod
    def _extract_author_from_content(doc_data: Dict) -> Optional[str]:
        """
        从文档内容中提取作者信息
        
        优先从"文档修改记录"表格中提取作者信息
        
        Args:
            doc_data: 解析后的文档数据
            
        Returns:
            作者信息（如果找到），否则返回None
        """
        # 查找"文档修改记录"表格
        # 表格标题可能包含：文档修改记录、修改记录、修订记录等
        modification_record_keywords = ["文档修改记录", "修改记录", "修订记录", "版本记录"]
        
        for table_data in doc_data.get("tables", []):
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])
            
            # 检查表格是否包含"作者"列
            if "作者" not in headers:
                continue
            
            # 查找"作者"列的索引
            author_col_idx = None
            for idx, header in enumerate(headers):
                if "作者" in header:
                    author_col_idx = idx
                break
            
            if author_col_idx is None:
                continue
            
            # 检查表格前的标题是否是"文档修改记录"相关
            # 查找表格在structured_content中的位置
            table_id = table_data.get("table_id", "")
            for item in doc_data.get("structured_content", []):
                if item.get("type") == "table" and item.get("table_id") == table_id:
                    # 查找表格前的标题
                    item_idx = doc_data.get("structured_content", []).index(item)
                    # 向前查找最近的标题
                    for prev_idx in range(item_idx - 1, max(-1, item_idx - 10), -1):
                        prev_item = doc_data.get("structured_content", [])[prev_idx]
                        if prev_item.get("type") == "heading" or prev_item.get("type") == "paragraph":
                            prev_text = prev_item.get("text", "")
                            # 检查是否包含修改记录相关的关键词
                            if any(keyword in prev_text for keyword in modification_record_keywords):
                                # 找到了"文档修改记录"表格，提取作者信息
                                # 从第一行数据中提取作者（通常是最新的版本）
                                if rows and len(rows) > 0:
                                    author = rows[0][author_col_idx].strip() if author_col_idx < len(rows[0]) else ""
                                    if author:
                                        # 去重：如果多行作者相同，只返回一个
                                        unique_authors = set()
                                        for row in rows:
                                            if author_col_idx < len(row):
                                                author_name = row[author_col_idx].strip()
                                                if author_name:
                                                    unique_authors.add(author_name)
                                        
                                        # 如果所有作者都相同，返回一个；否则返回所有作者（用顿号分隔）
                                        if len(unique_authors) == 1:
                                            return list(unique_authors)[0]
                                        elif len(unique_authors) > 1:
                                            return "、".join(sorted(unique_authors))
                                        else:
                                            return author
                                break
                    break
        
        # 如果没有找到，返回None（不显示作者）
        return None
    
    @staticmethod
    def _extract_overview_content(structured_content: List[Dict]) -> Optional[str]:
        """
        从结构化内容中提取功能概述
        
        查找第一个"概述"相关章节的内容，如功能概述、系统概述、项目概述等
        
        Args:
            structured_content: 结构化内容列表
            
        Returns:
            功能概述内容（如果找到），否则返回None
        """
        # 概述相关的关键词
        overview_keywords = ["概述", "功能说明", "功能描述", "系统说明", "项目说明", "简介", "背景"]
        
        # 查找包含概述关键词的章节
        for idx, item in enumerate(structured_content):
            if item.get("type") == "heading":
                heading_text = item.get("text", "").strip()
                # 检查标题是否包含概述关键词
                if any(keyword in heading_text for keyword in overview_keywords):
                    # 找到概述章节，收集该章节的内容
                    content_parts = []
                    current_level = item.get("level", 1)
                    
                    # 从下一个元素开始收集内容，直到遇到同级或更高级别的标题
                    for j in range(idx + 1, len(structured_content)):
                        next_item = structured_content[j]
                        if next_item.get("type") == "heading":
                            next_level = next_item.get("level", 1)
                            if next_level <= current_level:
                                # 遇到同级或更高级别的标题，停止收集
                                break
                        
                        # 收集段落内容
                        if next_item.get("type") == "paragraph":
                            text = next_item.get("text", "").strip()
                            if text:
                                content_parts.append(text)
                    
                    # 返回收集的内容（限制长度）
                    if content_parts:
                        full_content = "\n\n".join(content_parts)
                        # 限制最大长度为500字符
                        if len(full_content) > 500:
                            full_content = full_content[:500] + "..."
                        return full_content
        
        # 如果没有找到概述章节，尝试返回第一个段落内容作为概述
        for item in structured_content:
            if item.get("type") == "paragraph":
                text = item.get("text", "").strip()
                # 排除太短的段落（如标题行）
                if text and len(text) > 30:
                    if len(text) > 500:
                        return text[:500] + "..."
                    return text
        
        return None
    
    @staticmethod
    def _build_summary_content(doc_data: Dict, sections: List[Dict], document_id: str = None, upload_id: int = None, file_name: str = None) -> str:
        """
        构建总结文档（包含图片清单、表格清单、嵌入文档清单等）
        
        用于用户查看文档结构概览，不用于Graphiti处理
        
        Args:
            doc_data: 解析后的文档数据
            sections: 章节列表
            document_id: 文档ID（格式：upload_{upload_id}）
            upload_id: 上传ID（用于构建URL）
            file_name: 文件名（用于显示文档标题）
        """
        # 优先使用upload_id，如果没有则从document_id中提取
        if upload_id is None and document_id:
            if document_id.startswith("upload_"):
                try:
                    upload_id = int(document_id.replace("upload_", ""))
                except ValueError:
                    upload_id = None
        
        summary = "# 文档解析总结\n\n"
        
        # 文档基本信息
        metadata = doc_data.get("metadata", {})
        structured_content = doc_data.get("structured_content", [])
        
        # 尝试从文档内容中提取作者信息（优先从"文档修改记录"表格中提取）
        author_from_content = WordDocumentService._extract_author_from_content(doc_data)
        
        summary += "## 文档概览\n\n"
        
        # 基本信息
        summary += "### 基本信息\n\n"
        # 使用文件名作为文档标题（去掉扩展名）
        if file_name:
            # 去掉文件扩展名
            doc_title = file_name.rsplit('.', 1)[0] if '.' in file_name else file_name
        else:
            doc_title = metadata.get('title', '未命名文档')
        summary += f"- **文档标题**: {doc_title}\n"
        if author_from_content:
            summary += f"- **作者**: {author_from_content}\n"
        if metadata.get('created'):
            created = metadata.get('created')
            if isinstance(created, datetime):
                summary += f"- **创建时间**: {created.strftime('%Y-%m-%d %H:%M:%S')}\n"
            else:
                summary += f"- **创建时间**: {created}\n"
        if metadata.get('modified'):
            modified = metadata.get('modified')
            if isinstance(modified, datetime):
                summary += f"- **修改时间**: {modified.strftime('%Y-%m-%d %H:%M:%S')}\n"
            else:
                summary += f"- **修改时间**: {modified}\n"
        summary += "\n"
        
        # 章节结构（带层级的目录）
        summary += "### 章节结构\n\n"
        summary += "本文档包含以下主要章节：\n\n"
        for item in structured_content:
            if item.get("type") == "heading":
                level = item.get("level", 1)
                heading_text = item.get("text", "").strip()
                if heading_text:
                    indent = "  " * (level - 1)
                    summary += f"{indent}- {heading_text}\n"
        summary += "\n"
        
        # 功能概述（查找第一个"概述"相关章节的内容）
        overview_content = WordDocumentService._extract_overview_content(structured_content)
        if overview_content:
            summary += "### 功能概述\n\n"
            summary += overview_content + "\n\n"
        
        # 统计信息
        summary += "## 统计信息\n\n"
        summary += f"- **章节数**: {len(sections)}\n"
        summary += f"- **图片数**: {len(doc_data.get('images', []))}\n"
        summary += f"- **表格数**: {len(doc_data.get('tables', []))}\n"
        summary += f"- **链接数**: {len(doc_data.get('links', []))}\n"
        summary += f"- **嵌入文档数**: {len(doc_data.get('ole_objects', []))}\n"
        summary += f"- **文本长度**: {len(doc_data.get('text_content', ''))} 字符\n"
        summary += "\n"
        
        # 图片清单
        images = doc_data.get("images", [])
        if images:
            summary += "## 图片清单\n\n"
            for idx, image in enumerate(images, 1):
                image_id = image.get("image_id", f"image_{idx}")
                image_desc = image.get("description", "图片")
                section_title = image.get("section_title", "未知章节")
                relative_position = image.get("relative_position", 0.0)
                prev_context = image.get("prev_context", "")
                next_context = image.get("next_context", "")
                image_context = image.get("context", "")
                
                summary += f"### 图片 {idx} ({image_id})\n\n"
                summary += f"- **描述**: {image_desc}\n"
                summary += f"- **位置**: {section_title} (文档位置: {relative_position:.1%})\n"
                
                if upload_id:
                    image_url = f"/api/document-upload/{upload_id}/images/{image_id}"
                    summary += f"- **链接**: [查看图片]({image_url})\n"
                
                # 添加上下文信息
                summary += "\n**上下文信息**:\n"
                if prev_context:
                    summary += f"- **前文**: {prev_context}\n"
                if image_context:
                    summary += f"- **当前段落**: {image_context}\n"
                if next_context:
                    summary += f"- **后文**: {next_context}\n"
                if not prev_context and not image_context and not next_context:
                    summary += "- 无上下文信息\n"
                
                summary += "\n"
        
        # 表格清单（不包含上下文）
        tables = doc_data.get("tables", [])
        if tables:
            summary += "## 表格清单\n\n"
            for idx, table_data in enumerate(tables, 1):
                table_id = table_data.get("table_id", f"table_{idx}")
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                
                # 查找表格所在的章节
                section_title = "未知章节"
                structured_content = doc_data.get("structured_content", [])
                # 可以通过structured_content查找表格所在的章节
                for item_idx, item in enumerate(structured_content):
                    if item.get("type") == "table" and item.get("table_id") == table_id:
                        # 向前查找最近的章节标题
                        for i in range(item_idx, -1, -1):
                            prev_item = structured_content[i]
                            if prev_item.get("type") == "heading" and prev_item.get("level", 1) == 1:
                                section_title = prev_item.get("text", "未知章节")
                                break
                        break
                
                summary += f"### 表格 {idx} ({table_id})\n\n"
                summary += f"- **位置**: {section_title}\n"
                summary += f"- **行列数**: {len(rows)} 行 × {len(headers)} 列\n"
                
                # 内容摘要（前3行）
                if headers and rows:
                    summary += "\n**内容摘要**:\n\n"
                    # 表头
                    header_row = "| " + " | ".join(str(h) for h in headers) + " |\n"
                    separator = "| " + " | ".join(["---"] * len(headers)) + " |\n"
                    summary += header_row + separator
                    # 前3行数据
                    for row in rows[:3]:
                        row_str = "| " + " | ".join(str(cell) for cell in row) + " |\n"
                        summary += row_str
                    if len(rows) > 3:
                        summary += f"| ... (还有 {len(rows) - 3} 行) |\n"
                    summary += "\n"
                
                summary += "\n"
        
        # 嵌入文档清单
        ole_objects = doc_data.get("ole_objects", [])
        if ole_objects:
            summary += "## 嵌入文档清单\n\n"
            for idx, ole in enumerate(ole_objects, 1):
                ole_id = ole.get("ole_id", f"ole_{idx}")
                ole_name = ole.get("name", "嵌入文档")
                ole_type = ole.get("type", "未知类型")
                
                summary += f"### 嵌入文档 {idx} ({ole_id})\n\n"
                summary += f"- **名称**: {ole_name}\n"
                summary += f"- **类型**: {ole_type}\n"
                
                if upload_id:
                    preview_url = f"/api/document-upload/{upload_id}/ole/{ole_id}?view=preview"
                    download_url = f"/api/document-upload/{upload_id}/ole/{ole_id}?view=download"
                    summary += f"- **链接**: [查看]({preview_url}) | [下载]({download_url})\n"
                
                summary += "\n"
        
        # 链接清单（可选）
        links = doc_data.get("links", [])
        if links:
            summary += "## 链接清单\n\n"
            seen_links = set()
            for link in links:
                link_key = (link.get('url', ''), link.get('text', ''))
                if link_key not in seen_links:
                    seen_links.add(link_key)
                    url = link.get('url', '')
                    text = link.get('text', url)
                    link_type = link.get('type', 'external')
                    summary += f"- [{text}]({url}) ({link_type})\n"
            summary += "\n"
        
        return summary
    
    @staticmethod
    def _generate_document_summary(doc_data: Dict) -> str:
        """生成文档摘要"""
        summary = f"文档名称：{doc_data['metadata'].get('title', '')}\n"
        summary += f"作者：{doc_data['metadata'].get('author', '')}\n"
        created = doc_data['metadata'].get('created')
        if created:
            summary += f"创建时间：{created.strftime('%Y-%m-%d %H:%M:%S')}\n"
        summary += "\n文档概览：\n"
        summary += doc_data["text_content"][:2000]  # 前 2000 字符作为摘要
        return summary
    
    @staticmethod
    async def process_word_document(
        file_path: str,
        document_name: str,
        provider: str = "qianwen",
        max_tokens_per_section: int = 8000
    ) -> Dict[str, Any]:
        """
        处理 Word 文档，创建分层 Episode
        
        Args:
            file_path: Word 文档路径
            document_name: 文档名称
            provider: LLM 提供商
            max_tokens_per_section: 每个章节的最大 token 数
        
        Returns:
            处理结果，包含所有 Episode UUID
        """
        try:
            # Step 1: 解析 Word 文档
            logger.info(f"开始解析 Word 文档: {file_path}")
            doc_data = WordDocumentService._parse_word_document(file_path)
            logger.info(f"文档解析完成: {len(doc_data['structured_content'])} 个元素")
            
            # Step 2: 按章节分块
            sections = WordDocumentService._split_by_sections(
                doc_data["structured_content"],
                max_tokens=max_tokens_per_section
            )
            logger.info(f"文档分为 {len(sections)} 个章节")
            
            # Step 3: 获取 Graphiti 实例
            graphiti = get_graphiti_instance(provider)
            
            # 提取基础标识和版本号
            base_name = WordDocumentService._extract_base_name(document_name)
            version, version_number = WordDocumentService._extract_version(document_name)
            
            # 清理基础标识，只保留字母数字、破折号和下划线
            # Graphiti 要求 group_id 只能包含 alphanumeric characters, dashes, or underscores
            safe_base_name = WordDocumentService._sanitize_group_id(base_name)
            
            # 使用文档创建日期或当前日期
            doc_date = doc_data["metadata"].get("created")
            if doc_date and isinstance(doc_date, datetime):
                date_str = doc_date.strftime('%Y%m%d')
            else:
                date_str = datetime.now().strftime('%Y%m%d')
            
            # 生成基础 group_id（所有版本共享）
            group_id = f"doc_{safe_base_name}_{date_str}"
            logger.info(f"文档基础标识: {base_name}")
            logger.info(f"文档版本: {version} (版本号: {version_number})")
            logger.info(f"文档 group_id: {group_id}")
            
            # Step 4: 创建文档级 Episode（提取文档级别的实体）
            document_summary = WordDocumentService._generate_document_summary(doc_data)
            document_episode = await graphiti.add_episode(
                name=f"{document_name}_文档概览",
                episode_body=document_summary,
                source_description="Word文档",
                reference_time=doc_data["metadata"].get("created") or datetime.now(),
                entity_types={
                    "Requirement": ENTITY_TYPES.get("Requirement"),
                    "Document": ENTITY_TYPES.get("Document"),
                } if ENTITY_TYPES.get("Requirement") and ENTITY_TYPES.get("Document") else None,
                group_id=group_id
            )
            document_episode_uuid = document_episode.episode.uuid
            logger.info(f"文档级 Episode 创建完成: {document_episode_uuid}")
            
            # 更新文档级 Episode 的版本信息和文件路径
            from app.core.neo4j_client import neo4j_client
            update_version_query = """
            MATCH (e:Episodic)
            WHERE e.uuid = $episode_uuid
            SET e.version = $version,
                e.version_number = $version_number,
                e.document_name = $document_name,
                e.file_path = $file_path,
                e.original_filename = $original_filename
            RETURN e.uuid as uuid
            """
            neo4j_client.execute_write(update_version_query, {
                "episode_uuid": document_episode_uuid,
                "version": version,
                "version_number": version_number,
                "document_name": document_name,
                "file_path": file_path,
                "original_filename": os.path.basename(file_path)
            })
            logger.info(f"已更新文档级 Episode 版本信息和文件路径: version={version}, version_number={version_number}, file_path={file_path}")
            
            # 定义更新版本信息的函数（供后续使用）
            def update_episode_version(episode_uuid: str):
                neo4j_client.execute_write(update_version_query, {
                    "episode_uuid": episode_uuid,
                    "version": version,
                    "version_number": version_number,
                    "document_name": document_name,
                    "file_path": file_path,
                    "original_filename": os.path.basename(file_path)
                })
            
            # Step 5: 创建章节级 Episode
            section_episodes = []
            for idx, section in enumerate(sections):
                section_content = WordDocumentService._build_section_content(
                    section, doc_data, idx
                )
                
                section_episode = await graphiti.add_episode(
                    name=f"{document_name}_章节_{idx+1}_{section['title'][:20]}",
                    episode_body=section_content,
                    source_description="Word文档章节",
                    reference_time=doc_data["metadata"].get("created") or datetime.now(),
                    entity_types=ENTITY_TYPES,
                    edge_types=EDGE_TYPES,
                    edge_type_map=EDGE_TYPE_MAP,
                    group_id=group_id,
                    previous_episode_uuids=[document_episode_uuid]
                )
                
                section_episode_uuid = section_episode.episode.uuid
                section_episodes.append(section_episode_uuid)
                
                # 更新章节级 Episode 的版本信息
                update_episode_version(section_episode_uuid)
                
                logger.info(f"章节 {idx+1} Episode 创建完成: {section_episode_uuid}")
            
            # Step 6: 处理图片，为每张图片创建独立的Episode
            image_episodes = []
            if doc_data["images"]:
                logger.info(f"开始处理 {len(doc_data['images'])} 张图片")
                for idx, image in enumerate(doc_data["images"]):
                    image_id = image.get("image_id", f"image_{idx+1}")
                    image_desc = image.get("description", f"图片 {idx+1}")
                    image_context = image.get("context", "")
                    image_url = f"/api/word-document/{group_id}/images/{image_id}"
                    
                    # 获取增强的上下文信息
                    prev_context = image.get('prev_context', '')
                    next_context = image.get('next_context', '')
                    section_title = image.get('section_title', '')
                    relative_position = image.get('relative_position', 0.0)
                    match_method = image.get('match_method', 'unknown')
                    match_confidence = image.get('match_confidence', 0.0)
                    file_size = image.get('file_size', 0)
                    file_format = image.get('file_format', 'UNKNOWN')
                    
                    # 构建图片Episode的内容（增强版：包含更多元数据和上下文）
                    image_content = f"""## 图片信息

**图片ID**: {image_id}
**描述**: {image_desc}
**文件路径**: {image.get('file_path', '')}
**相对路径**: {image.get('relative_path', '')}
**文件大小**: {file_size} 字节
**文件格式**: {file_format}
**匹配方法**: {match_method}
**匹配置信度**: {match_confidence:.2f}
**文档位置**: {relative_position:.1%}
"""
                    
                    # 添加章节信息
                    if section_title:
                        image_content += f"**所属章节**: {section_title}\n\n"
                    
                    # 添加完整的上下文信息
                    image_content += "### 上下文信息\n\n"
                    if prev_context:
                        image_content += f"**前文**: {prev_context}\n\n"
                    if image_context:
                        image_content += f"**当前段落**: {image_context}\n\n"
                    if next_context:
                        image_content += f"**后文**: {next_context}\n\n"
                    if not prev_context and not image_context and not next_context:
                        image_content += "无上下文信息\n\n"
                    
                    image_content += f"""### 图片链接
![{image_desc}]({image_url})

### 图片说明
这是一张从Word文档中提取的图片，位于文档的相应位置（位置: {relative_position:.1%}）。图片可能包含流程图、示意图、图表或其他可视化内容。

**匹配信息**: 通过{match_method}方法匹配，置信度为{match_confidence:.0%}。
"""
                    
                    # 创建图片Episode
                    image_episode = await graphiti.add_episode(
                        name=f"{document_name}_图片_{idx+1}_{image_desc[:20]}",
                        episode_body=image_content,
                        source_description="Word文档图片",
                        reference_time=doc_data["metadata"].get("created") or datetime.now(),
                        entity_types=ENTITY_TYPES,
                        edge_types=EDGE_TYPES,
                        edge_type_map=EDGE_TYPE_MAP,
                        group_id=group_id,
                        previous_episode_uuids=[document_episode_uuid]
                    )
                    
                    image_episode_uuid = image_episode.episode.uuid
                    image_episodes.append(image_episode_uuid)
                    
                    # 更新图片 Episode 的版本信息
                    update_episode_version(image_episode_uuid)
                    
                    logger.info(f"图片 {idx+1} Episode 创建完成: {image_episode_uuid}")
            
            # Step 7: 处理表格，为每个表格创建独立的Episode
            table_episodes = []
            if doc_data["tables"]:
                logger.info(f"开始处理 {len(doc_data['tables'])} 个表格")
                for idx, table_data in enumerate(doc_data["tables"]):
                    # 格式化表格为标准Markdown格式（用于Episode内容）
                    table_markdown = WordDocumentService._format_table_as_markdown(table_data)
                    
                    # 构建表格Episode的内容
                    table_content = f"""## 表格信息

**表格序号**: {idx+1}
**表格ID**: {table_data.get('table_id', f'table_{idx+1}')}
**行数**: {len(table_data.get('rows', []))}
**列数**: {len(table_data.get('headers', []))}

### 表格内容

{table_markdown}

### 表格说明
这是从Word文档中提取的表格数据，使用标准Markdown表格格式，包含结构化的信息。
"""
                    
                    # 创建表格Episode
                    table_episode = await graphiti.add_episode(
                        name=f"{document_name}_表格_{idx+1}",
                        episode_body=table_content,
                        source_description="Word文档表格",
                        reference_time=doc_data["metadata"].get("created") or datetime.now(),
                        entity_types=ENTITY_TYPES,
                        edge_types=EDGE_TYPES,
                        edge_type_map=EDGE_TYPE_MAP,
                        group_id=group_id,
                        previous_episode_uuids=[document_episode_uuid]
                    )
                    
                    table_episode_uuid = table_episode.episode.uuid
                    table_episodes.append(table_episode_uuid)
                    
                    # 更新表格 Episode 的版本信息
                    update_episode_version(table_episode_uuid)
                    
                    logger.info(f"表格 {idx+1} Episode 创建完成: {table_episode_uuid}")
            
            return {
                "success": True,
                "document_id": group_id,
                "document_name": document_name,
                "document_episode_uuid": document_episode_uuid,
                "section_episodes": section_episodes,
                "image_episodes": image_episodes,
                "table_episodes": table_episodes,
                "statistics": {
                    "total_sections": len(sections),
                    "total_images": len(doc_data["images"]),
                    "total_tables": len(doc_data["tables"]),
                    "total_links": len(doc_data["links"])
                }
            }
        except Exception as e:
            logger.error(f"处理 Word 文档失败: {e}", exc_info=True)
            raise


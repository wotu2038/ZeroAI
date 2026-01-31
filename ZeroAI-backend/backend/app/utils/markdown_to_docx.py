"""
Markdown 转 DOCX 工具
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Optional
import logging

logger = logging.getLogger(__name__)


def markdown_to_docx(markdown_text: str, output_path: Optional[str] = None) -> bytes:
    """
    将 Markdown 文本转换为 DOCX 文档
    
    Args:
        markdown_text: Markdown 格式的文本
        output_path: 输出文件路径（可选，如果提供则保存到文件）
    
    Returns:
        DOCX 文件的字节内容
    """
    doc = Document()
    
    # 设置中文字体
    def set_chinese_font(run):
        from docx.oxml.ns import qn as _qn
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(_qn('w:eastAsia'), '宋体')
    
    # 设置段落字体
    def set_paragraph_font(paragraph, font_name='宋体', font_size=12):
        from docx.oxml.ns import qn as _qn
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(_qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
    
    # 按行处理
    lines = markdown_text.split('\n')
    i = 0
    
    # 检测ASCII艺术图（包含大量特殊字符的行）
    def is_ascii_art(line):
        """检测是否是ASCII艺术图"""
        if not line:
            return False
        # 如果包含大量特殊字符（+、-、|、/、\等），可能是ASCII艺术图
        special_chars = set('+-|/\\<>[]{}()=')
        special_count = sum(1 for c in line if c in special_chars)
        # 如果特殊字符占比超过20%，或者包含多个连续的+、-、|，可能是ASCII艺术图
        has_box_chars = any(pattern in line for pattern in ['+--', '---', '|||', '+++'])
        return special_count > len(line) * 0.2 or has_box_chars  # 20%以上是特殊字符或包含框线字符
    
    while i < len(lines):
        line = lines[i]
        line_stripped = line.strip()
        
        # 跳过空行
        if not line_stripped:
            i += 1
            continue
        
        # 标题处理
        if line_stripped.startswith('#'):
            level = len(line_stripped) - len(line_stripped.lstrip('#'))
            title_text = line_stripped.lstrip('#').strip()
            
            if title_text:
                heading = doc.add_heading(title_text, level=min(level, 9))
                set_paragraph_font(heading, font_name='黑体', font_size=16 - level)
                i += 1
                continue
        
        # ASCII艺术图处理（在代码块之前检测）
        if is_ascii_art(line_stripped):
            # 收集连续的ASCII艺术图行（包括空行，保持格式）
            ascii_lines = []
            j = i
            while j < len(lines):
                current_line = lines[j]
                current_stripped = current_line.strip()
                # 如果是ASCII艺术图行或空行，继续收集
                if is_ascii_art(current_stripped) or not current_stripped:
                    ascii_lines.append(current_line)  # 保留原始行，包括前导空格
                    j += 1
                else:
                    # 如果遇到非ASCII艺术图行，停止收集
                    break
            
            if ascii_lines:
                # 使用等宽字体显示ASCII艺术图
                ascii_text = '\n'.join(ascii_lines)
                p = doc.add_paragraph(ascii_text)
                p.style = 'No Spacing'
                # 设置等宽字体
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                # 保持原始格式（不自动换行）
                p.paragraph_format.keep_together = True
                # 设置段落格式为不换行（保持原始布局）
                try:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import qn as _qn
                    # 设置不自动换行
                    pPr = p._element.get_or_add_pPr()
                    noWrap = parse_xml(f'<w:noWrap xmlns:w="{_qn("w").split("}")[0][1:]}" />')
                    pPr.append(noWrap)
                except:
                    pass  # 如果设置失败，至少保持等宽字体
                i = j
                continue
        
        # 列表处理
        if line_stripped.startswith('- ') or line_stripped.startswith('* ') or re.match(r'^\d+\.\s', line_stripped):
            # 无序列表
            if line_stripped.startswith('- ') or line_stripped.startswith('* '):
                list_text = line_stripped[2:].strip()
                p = doc.add_paragraph(list_text, style='List Bullet')
            # 有序列表
            else:
                list_text = re.sub(r'^\d+\.\s', '', line_stripped).strip()
                p = doc.add_paragraph(list_text, style='List Number')
            
            set_paragraph_font(p)
            i += 1
            continue
        
        # 表格处理（简单表格）
        if '|' in line_stripped and i + 1 < len(lines):
            # 检查下一行是否是分隔符
            next_line = lines[i + 1].strip() if i + 1 < len(lines) else ''
            if re.match(r'^[\|\s\-:]+$', next_line):
                # 这是一个表格
                table_lines = []
                j = i
                while j < len(lines) and '|' in lines[j]:
                    table_lines.append(lines[j])
                    j += 1
                
                if len(table_lines) >= 2:
                    # 解析表格
                    headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                    if headers:
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Light Grid Accent 1'
                        
                        # 设置表头
                        header_cells = table.rows[0].cells
                        for idx, header in enumerate(headers):
                            header_cells[idx].text = header
                            set_paragraph_font(header_cells[idx].paragraphs[0], font_name='黑体')
                        
                        # 添加数据行
                        for row_line in table_lines[2:]:  # 跳过表头和分隔符
                            cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
                            if len(cells) == len(headers):
                                row = table.add_row()
                                for idx, cell in enumerate(cells):
                                    row.cells[idx].text = cell
                                    set_paragraph_font(row.cells[idx].paragraphs[0])
                        
                        i = j
                        continue
        
        # 代码块处理
        if line_stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                # 设置代码样式（等宽字体）
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                # 保持原始格式
                p.paragraph_format.keep_together = True
            i += 1
            continue
        
        # 链接处理 [text](url)
        link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
        if re.search(link_pattern, line_stripped):
            p = doc.add_paragraph()
            last_pos = 0
            for match in re.finditer(link_pattern, line_stripped):
                # 添加链接前的文本
                if match.start() > last_pos:
                    text_before = line_stripped[last_pos:match.start()]
                    run = p.add_run(text_before)
                    set_chinese_font(run)
                
                # 添加链接
                link_text = match.group(1)
                link_url = match.group(2)
                run = p.add_run(link_text)
                set_chinese_font(run)
                run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                run.underline = True
                # 添加超链接
                try:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls, qn
                    hyperlink = parse_xml(
                        f'<w:hyperlink r:id="{link_url}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                    )
                    run._element.getparent().replace(run._element, hyperlink)
                except:
                    pass  # 如果超链接添加失败，至少保留文本
                
                last_pos = match.end()
            
            # 添加链接后的文本
            if last_pos < len(line_stripped):
                text_after = line_stripped[last_pos:]
                run = p.add_run(text_after)
                set_chinese_font(run)
            
            set_paragraph_font(p)
            i += 1
            continue
        
        # 普通段落处理（处理加粗、斜体等格式）
        # 保留原始行的内容（包括前导空格），但处理格式
        p = doc.add_paragraph()
        process_formatted_text(p, line_stripped)
        set_paragraph_font(p)
        i += 1
    
    # 如果没有内容，至少添加一个段落
    if len(doc.paragraphs) == 0:
        doc.add_paragraph(markdown_text)
    
    # 保存或返回
    if output_path:
        doc.save(output_path)
        with open(output_path, 'rb') as f:
            return f.read()
    else:
        import io
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()


def process_formatted_text(paragraph, text: str):
    """
    处理带格式的文本（加粗、斜体等）
    """
    # 处理加粗 **text** 或 __text__
    text = re.sub(r'\*\*([^\*]+)\*\*', r'<bold>\1</bold>', text)
    text = re.sub(r'__([^_]+)__', r'<bold>\1</bold>', text)
    
    # 处理斜体 *text* 或 _text_
    text = re.sub(r'(?<!\*)\*([^\*]+)\*(?!\*)', r'<italic>\1</italic>', text)
    text = re.sub(r'(?<!_)_([^_]+)_(?!_)', r'<italic>\1</italic>', text)
    
    # 处理删除线 ~~text~~
    text = re.sub(r'~~([^~]+)~~', r'<strike>\1</strike>', text)
    
    # 处理行内代码 `code`
    text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
    
    # 分割文本并添加格式
    parts = re.split(r'(<[^>]+>[^<]+</[^>]+>)', text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('<bold>'):
            run_text = part.replace('<bold>', '').replace('</bold>', '')
            run = paragraph.add_run(run_text)
            run.bold = True
        elif part.startswith('<italic>'):
            run_text = part.replace('<italic>', '').replace('</italic>', '')
            run = paragraph.add_run(run_text)
            run.italic = True
        elif part.startswith('<strike>'):
            run_text = part.replace('<strike>', '').replace('</strike>', '')
            run = paragraph.add_run(run_text)
            run.font.strike = True
        elif part.startswith('<code>'):
            run_text = part.replace('<code>', '').replace('</code>', '')
            run = paragraph.add_run(run_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
        
        # 设置中文字体
        if hasattr(run, '_element'):
            from docx.oxml.ns import qn as _qn
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(_qn('w:eastAsia'), '宋体')

